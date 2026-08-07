"""
==============================================================================
MULTI TOOL AGENT — LANGGRAPH WORKSTATION (CLI VERSION)
==============================================================================
Production-grade AI workstation CLI with:
- Multi-modal ingestion (PDF, DOCX, TXT, CSV, Excel, Images incl. GIF, URLs, MySQL)
- LangGraph agent tool routing with checkpointer memory
- Hybrid retrieval (BM25 + FAISS vector search)
- Query optimization (NONE, EXPANSION, DECOMPOSITION, HYDE)
- Auto-summarization & 7-day chat history persistence (atomic file writes)
- Self-healing code execution for data analysis (hardened sandbox wrapper)
- DuckDuckGo live web search integration
- Robust SQL safety net (validated table refs + LIMIT skipping for aggregations)
- LangSmith observability (opt-in via env vars; zero overhead when disabled)
==============================================================================
"""

import os
import re
import ast
import sys
import uuid
import json
import base64
import tempfile
import time
import threading
from functools import lru_cache
from datetime import datetime, timedelta
from typing import List, Optional, Dict, Any, Union, Tuple, Annotated, TypedDict

# Data processing, database connectivity, and environment management libraries
import pandas as pd
import requests
import pymysql
from dotenv import load_dotenv
from sqlalchemy import create_engine, text
from urllib.parse import quote_plus

# LangChain core primitives for structured messaging, doc representations, and runnables
from langchain_core.messages import (
    BaseMessage, AIMessage, HumanMessage, SystemMessage,
)
from langchain_core.documents import Document
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnableLambda
from langchain_core.tools import tool

# Integration wrappers for Google Generative AI models (Chat & Embeddings)
from langchain_google_genai import (
    ChatGoogleGenerativeAI,
    GoogleGenerativeAIEmbeddings,
)

# Text chunking, Web scraping, PyPDF parsing, Vector index, and Sparse BM25 retrieval
from langchain_community.document_loaders import WebBaseLoader, PyPDFLoader
from langchain_community.retrievers import BM25Retriever
try:
    from langchain_classic.retrievers.ensemble import EnsembleRetriever
except ImportError:
    from langchain_community.retrievers import EnsembleRetriever

from langchain_community.vectorstores import FAISS
from langchain_ollama import OllamaEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

# LangGraph state management, graph orchestration, and in-memory checkpointing
from langgraph.graph import StateGraph, START, END
from langgraph.graph.message import add_messages
from langgraph.prebuilt import ToolNode, tools_condition
from langgraph.checkpoint.memory import InMemorySaver

# High-fidelity document converter (Docling) for rich markdown layout preservation
try:
    from docling.document_converter import DocumentConverter
    DOCLING_AVAILABLE = True
except ImportError:
    DOCLING_AVAILABLE = False

# DuckDuckGo live web search engine integration
try:
    from duckduckgo_search import DDGS
    DDG_AVAILABLE = True
except ImportError:
    DDG_AVAILABLE = False

# LangSmith Observability Tracing & Client setup (with graceful fallback if not installed)
try:
    from langchain_core.tracers import LangChainTracer
    from langsmith import Client as LangSmithClient
    LANGSMITH_AVAILABLE = True
except ImportError:
    LANGSMITH_AVAILABLE = False
    LangChainTracer = None
    LangSmithClient = None

# Load environment variables from local .env file if present
load_dotenv()


# ================================================================================
# CONFIGURATION CLASS
# Centralized application configuration parameters, system thresholds, model names,
# and sandbox execution bounds.
# ================================================================================

class AppConfig:
    """
    Central repository for application-wide defaults, environment configurations,
    retrieval chunking thresholds, model identifiers, and sandbox security parameters.
    """
    # Active AI Model Selections
    CHAT_MODEL: str = os.getenv("CHAT_MODEL", "gemini-2.5-flash")
    VISION_MODEL: str = os.getenv("VISION_MODEL", "gemini-2.5-flash")
    EMBED_MODEL: str = os.getenv("EMBED_MODEL", "nomic-embed-text")
    GOOGLE_EMBED_MODEL: str = os.getenv("GOOGLE_EMBED_MODEL", "models/text-embedding-004")
    GOOGLE_API_KEY: str = os.getenv("GOOGLE_API_KEY", "").strip()

    # Dynamic Chunking Dimensions (Small: <=5 pgs, Medium: <=15 pgs, Large: >15 pgs)
    DEFAULT_CHUNK_SIZE_SMALL: int = 500
    DEFAULT_CHUNK_OVERLAP_SMALL: int = 50
    DEFAULT_CHUNK_SIZE_MEDIUM: int = 1000
    DEFAULT_CHUNK_OVERLAP_MEDIUM: int = 100
    DEFAULT_CHUNK_SIZE_LARGE: int = 2000
    DEFAULT_CHUNK_OVERLAP_LARGE: int = 200

    # Retrieval Configuration Parameters
    RETRIEVER_K: int = 4
    ENSEMBLE_WEIGHTS: List[float] = [0.3, 0.7]  # [Sparse BM25 Weight, Dense FAISS Weight]

    # Execution Bounds & Timeouts
    OLLAMA_TIMEOUT: float = float(os.getenv("OLLAMA_TIMEOUT", "5.0"))
    MAX_CODE_EXECUTION_ATTEMPTS: int = 3  # Maximum Reflexion cycles for pandas code generation
    MIN_TEXT_LENGTH_FOR_PROCESSING: int = 10
    MAX_FILE_SIZE_MB: float = 50.0

    # Chat History & Persistence Parameters
    HISTORY_FILE: str = "chat_history_db.json"
    HISTORY_RETENTION_DAYS: int = 7

    # Summarization Thresholds
    MAX_VERBATIM_MESSAGES: int = 6
    SUMMARY_TRIGGER_LENGTH: int = 10
    MAX_SUMMARY_TOKENS: int = 400

    # System Guardrails
    MAX_PROMPT_CHARS: int = 200_000
    RECURSION_LIMIT: int = 12

    # Restricted builtins dictionary for sandboxed Python code execution
    SANDBOX_BUILTINS: dict = {
        "len": len, "sum": sum, "min": min, "max": max,
        "abs": abs, "round": round, "list": list, "dict": dict,
        "set": set, "tuple": tuple, "str": str, "int": int,
        "float": float, "bool": bool, "range": range,
        "any": any, "all": all, "sorted": sorted, "zip": zip,
        "enumerate": enumerate,
    }

    # LangSmith Telemetry Configuration
    LANGCHAIN_TRACING_V2: bool = os.getenv("LANGCHAIN_TRACING_V2", "false").lower() == "true"
    LANGCHAIN_API_KEY: str = os.getenv("LANGCHAIN_API_KEY", "").strip()
    LANGCHAIN_PROJECT: str = os.getenv("LANGCHAIN_PROJECT", "multi-tool-agent")
    LANGCHAIN_ENDPOINT: str = os.getenv(
        "LANGCHAIN_ENDPOINT", "https://api.smith.langchain.com"
    )


# ================================================================================
# LANGSMITH OBSERVABILITY HELPERS
# Infrastructure for end-to-end tracing, latency monitoring, and telemetry logging.
# ================================================================================

def is_langsmith_enabled() -> bool:
    """
    Determines whether LangSmith tracing is fully configured and operational.
    Requires: langsmith package installed, tracing env var enabled, and a valid API key.
    """
    return (
        LANGSMITH_AVAILABLE
        and AppConfig.LANGCHAIN_TRACING_V2
        and bool(AppConfig.LANGCHAIN_API_KEY)
    )


def get_langsmith_callbacks() -> List:
    """
    Constructs the list of LangChain callback handlers for explicit tracing.
    Returns an empty list if LangSmith is not enabled (fails silently to prevent pipeline interruption).
    """
    if not is_langsmith_enabled():
        return []
    callbacks = []
    try:
        client = LangSmithClient(
            api_url=AppConfig.LANGCHAIN_ENDPOINT,
            api_key=AppConfig.LANGCHAIN_API_KEY,
        )
        tracer = LangChainTracer(
            project_name=AppConfig.LANGCHAIN_PROJECT,
            client=client,
        )
        callbacks.append(tracer)
    except Exception as e:
        print(f"⚠️ LangSmith tracer initialization failed: {e}")
    return callbacks


def get_langsmith_config(
    run_name: Optional[str] = None,
    tags: Optional[List[str]] = None,
    metadata: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """
    Builds a standard Runnable config dictionary carrying LangSmith callbacks and metadata tags.
    Pass this dictionary as the `config=` argument to any chain or agent invocation.
    """
    config: Dict[str, Any] = {}
    callbacks = get_langsmith_callbacks()
    if callbacks:
        config["callbacks"] = callbacks
    if run_name:
        config["run_name"] = run_name
    if tags:
        config["tags"] = tags
    if metadata:
        config["metadata"] = metadata
    return config


# ================================================================================
# HARDENED SANDBOX: SAFE PROXY
# Defensive proxy object designed to neutralize Python Method Resolution Order (MRO)
# introspection and sandbox escapes during dynamically generated code execution.
# ================================================================================

class _SafeProxy:
    """
    Wraps objects (such as pandas DataFrames, Series, or modules) exposed to user code.
    Interprets dunder access through a minimal whitelist, blocking access to dangerous
    attributes like `__class__`, `__mro__`, or `__subclasses__` used in object breakout attacks.
    """

    # Whitelist of safe dunder methods required for basic data manipulation and arithmetic
    _ALLOWED_DUNDERS = frozenset({
        "__getitem__", "__setitem__", "__delitem__",
        "__iter__", "__next__", "__len__", "__contains__",
        "__add__", "__sub__", "__mul__", "__truediv__", "__floordiv__", "__mod__",
        "__radd__", "__rsub__", "__rmul__", "__rtruediv__",
        "__neg__", "__pos__", "__abs__",
        "__eq__", "__ne__", "__lt__", "__le__", "__gt__", "__ge__",
        "__hash__", "__bool__",
        "__repr__", "__str__",
        "__getattr__",
    })

    _BLOCKED_DUNDER_PREFIXES = ("_",)

    def __init__(self, obj):
        object.__setattr__(self, "_obj", obj)

    def __getattr__(self, name):
        # Intercept double-underscore dunder access
        if name.startswith("__") and name.endswith("__"):
            if name in self._ALLOWED_DUNDERS:
                return getattr(object.__getattribute__(self, "_obj"), name)
            raise AttributeError(
                f"Access to dunder '{name}' is blocked for security."
            )
        return getattr(object.__getattribute__(self, "_obj"), name)

    def __setattr__(self, name, value):
        if name.startswith("__") and name.endswith("__") and name != "_obj":
            raise AttributeError("Cannot assign dunder attributes.")
        object.__setattr__(name, value)

    def __getitem__(self, key):
        return object.__getattribute__(self, "_obj")[key]

    def __iter__(self):
        return iter(object.__getattribute__(self, "_obj"))

    def __len__(self):
        return len(object.__getattribute__(self, "_obj"))

    def __contains__(self, item):
        return item in object.__getattribute__(self, "_obj")

    def __repr__(self):
        return repr(object.__getattribute__(self, "_obj"))


# ================================================================================
# HARDENED SANDBOX: AST STATIC VALIDATOR
# Static analysis pass over LLM-generated code prior to evaluation via `exec()`.
# ================================================================================

# Blacklisted module names, builtins, and dangerous invocation names
_SANDBOX_FORBIDDEN_NAMES = frozenset({
    "eval", "exec", "compile", "open", "__import__", "globals", "locals",
    "vars", "getattr", "setattr", "delattr", "input", "help", "memoryview",
    "os", "sys", "subprocess", "shutil", "socket", "ctypes", "importlib",
    "pathlib", "pickle", "builtins",
})

# Whitelisted dunder attributes allowed during AST traversal
_SANDBOX_ALLOWED_DUNDER_ATTRS = frozenset({
    "__len__", "__iter__", "__next__", "__contains__",
    "__add__", "__sub__", "__mul__", "__truediv__", "__floordiv__", "__mod__",
    "__radd__", "__rsub__", "__rmul__", "__rtruediv__",
    "__neg__", "__pos__", "__abs__",
    "__eq__", "__ne__", "__lt__", "__le__", "__gt__", "__ge__",
    "__repr__", "__str__",
})


def validate_sandbox_code(code: str) -> Optional[str]:
    """
    Parses dynamic code into an Abstract Syntax Tree (AST) and scans every node for forbidden
    imports, dangerous builtin calls, or unapproved dunder attributes.
    Returns None if safe, or a specific violation description string if unsafe.
    """
    try:
        tree = ast.parse(code, mode="exec")
    except SyntaxError as e:
        return f"Generated code has a syntax error: {e}"

    # Walk the AST to inspect code structures statically
    for node in ast.walk(tree):
        # Reject explicit import statements
        if isinstance(node, (ast.Import, ast.ImportFrom)):
            return "Import statements are not allowed in the sandbox."
        
        # Inspect attribute access (e.g., obj.__class__)
        if isinstance(node, ast.Attribute):
            attr = node.attr
            if attr.startswith("__") and attr.endswith("__"):
                if attr not in _SANDBOX_ALLOWED_DUNDER_ATTRS:
                    return (
                        f"Access to dunder attribute '{attr}' is blocked "
                        "(sandbox-escape primitive)."
                    )
        
        # Inspect direct identifier calls and variable names
        if isinstance(node, ast.Name) and node.id in _SANDBOX_FORBIDDEN_NAMES:
            return f"Reference to '{node.id}' is not allowed in the sandbox."

    return None


# ================================================================================
# LLM / EMBEDDING CACHING & PROVIDERS
# Efficient initialization of Google Gemini models and hybrid local/cloud embeddings.
# ================================================================================

@lru_cache(maxsize=16)
def _get_cached_chat_llm(model_name: str, temperature: float = 0.0):
    """
    Retrieves or instantiates a cached ChatGoogleGenerativeAI client using LRU memory caching.
    Prevents redundant API client re-creation across conversational turns.
    """
    if not AppConfig.GOOGLE_API_KEY:
        print("❌ Error: GOOGLE_API_KEY is missing! Set it in .env or environment variables.")
        sys.exit(1)
    return ChatGoogleGenerativeAI(
        model=model_name,
        google_api_key=AppConfig.GOOGLE_API_KEY,
        temperature=temperature,
        max_retries=2,
        timeout=60,
    )


class _EmbeddingProvider:
    """
    Adaptive embedding provider that attempts local Ollama server connectivity first
    and gracefully falls back to Google Cloud Embeddings if Ollama is unreachable.
    """
    def __init__(self):
        self.mode = "google"
        try:
            # Check for active local Ollama daemon
            r = requests.get(
                "http://localhost:11434/api/tags",
                timeout=AppConfig.OLLAMA_TIMEOUT,
            )
            if r.status_code == 200:
                installed = {
                    m.get("name", "").split(":")[0]
                    for m in r.json().get("models", [])
                }
                if AppConfig.EMBED_MODEL.split(":")[0] in installed:
                    self._client = OllamaEmbeddings(model=AppConfig.EMBED_MODEL)
                    self.mode = "ollama"
                    return
        except requests.RequestException:
            pass  # Local Ollama unavailable; proceed to Google fallback

        # Fall back to Google Generative AI Cloud Embeddings
        if not AppConfig.GOOGLE_API_KEY:
            raise ValueError("GOOGLE_API_KEY required for fallback embeddings.")
        self._client = GoogleGenerativeAIEmbeddings(
            model=AppConfig.GOOGLE_EMBED_MODEL,
            google_api_key=AppConfig.GOOGLE_API_KEY,
        )

    @property
    def client(self):
        return self._client


@lru_cache(maxsize=1)
def _get_cached_embeddings():
    """Singleton getter for the adaptive embedding provider."""
    return _EmbeddingProvider()


def get_embedding_provider():
    """Returns the initialized embedding model client."""
    return _get_cached_embeddings().client


# ================================================================================
# UTILITIES & HELPER FUNCTIONS
# Input validation, error sanitization, exponential backoff execution, and formatting.
# ================================================================================

def extract_message_content(msg) -> str:
    """
    Normalizes complex multi-modal or structured message payloads into a clean string format.
    Handles raw strings, list-of-dicts (Gemini block layout), and direct content attributes.
    """
    if msg is None:
        return ""
    content = msg.content if hasattr(msg, "content") else msg

    if isinstance(content, str):
        return content

    if isinstance(content, list):
        text_parts = []
        for block in content:
            if isinstance(block, dict):
                if "text" in block and block["text"]:
                    text_parts.append(block["text"])
                elif block.get("type") == "text" and block.get("text"):
                    text_parts.append(block["text"])
            elif isinstance(block, str):
                text_parts.append(block)
        return "\n".join(text_parts) if text_parts else ""

    if isinstance(content, dict):
        if content.get("type") == "text" and content.get("text"):
            return content["text"]
        if "text" in content and content["text"]:
            return content["text"]
        return str(content)

    return str(content) if content is not None else ""


def invoke_with_retry(chain, inputs, max_retries: int = 3, config: Optional[Dict[str, Any]] = None):
    """
    Executes a LangChain runnable or graph node with exponential backoff retries for transient
    network errors, API rate limits (HTTP 429), or temporary server unavailability.
    """
    last_exception = None
    for attempt in range(max_retries):
        try:
            if isinstance(inputs, list):
                return chain.invoke(inputs, config=config) if config else chain.invoke(inputs)
            if isinstance(inputs, str):
                return chain.invoke(inputs, config=config) if config else chain.invoke(inputs)
            if isinstance(inputs, dict):
                try:
                    return chain.invoke(inputs, config=config) if config else chain.invoke(inputs)
                except TypeError:
                    if "messages" in inputs and isinstance(inputs["messages"], list):
                        return chain.invoke(inputs["messages"], config=config) if config else chain.invoke(inputs["messages"])
                    if "input" in inputs:
                        return chain.invoke(inputs["input"], config=config) if config else chain.invoke(inputs["input"])
                    raise
            return chain.invoke(inputs, config=config) if config else chain.invoke(inputs)
        except Exception as e:
            last_exception = e
            err_lower = str(e).lower()
            is_retryable = any(
                kw in err_lower
                for kw in ("429", "rate", "quota", "timeout", "connection",
                           "temporary", "unavailable")
            )
            if is_retryable and attempt < max_retries - 1:
                time.sleep(2 ** attempt)  # Exponential sleep: 1s, 2s, 4s...
                continue
            if attempt == max_retries - 1:
                raise RuntimeError(f"Failed after {max_retries} retries: {e}") from e
            raise
    raise RuntimeError("invoke_with_retry exited unexpectedly") from last_exception


def friendly_error(error: Exception) -> str:
    """
    Converts raw technical stack traces into human-readable user error messages with action guidance.
    """
    normalized = str(error).lower().replace(" ", "_").replace("-", "_")
    err_short = str(error)[:500]

    error_map = [
        (r"(?i)\bapi[_-]?key[_-]?invalid\b", "🔑 Invalid API key. Check GOOGLE_API_KEY in .env."),
        (r"(?i)\bpermission[_-]?denied\b", "🔑 Permission denied for this API key."),
        (r"(?i)\bquota[_-]?exceeded\b", "⏱️ API quota exceeded. Please try again later."),
        (r"(?i)\b429\b|\brate[_-]?limit\b", "⏱️ Rate limit hit. Slowing down requests."),
        (r"(?i)\bfile[_-]?too[_-]?large\b", "📦 File exceeds size threshold (max 50MB)."),
        (r"(?i)\bunsupported[_-]?format\b", "📄 Format not supported. Use PDF, DOCX, TXT, CSV, Excel, or Images."),
        (r"(?i)(?<!\w)memory\b|(?i)\boom\b|(?i)out[_\s]?of[_\s]?memory", "💾 Out of memory. Try a smaller dataset or restart the app."),
        (r"(?i)\btimeout\b", "⏱️ Request timed out. Try reducing input size."),
        (r"(?i)\bconnection\b", "🌐 Network connectivity issue."),
        (r"(?i)\bmysql\b|py.?mysql", "🗄️ MySQL connection error. Verify database credentials."),
        (r"(?i)\b404\b|(?i)\bnot[_-]?found\b", "🔍 Resource not found. Check the URL or table name."),
        (r"(?i)\bsandbox\b", "🛡️ Sandbox refused unsafe operation."),
    ]
    for pattern, message in error_map:
        if re.search(pattern, normalized):
            return message
    return f"❌ Error: {err_short}"


def validate_file_size(file_path: str, max_mb: float = 50.0) -> Tuple[bool, str]:
    """Validates that a local target file does not exceed the allowed maximum megabyte threshold."""
    if not os.path.exists(file_path):
        return False, f"File does not exist: {file_path}"
    size_mb = os.path.getsize(file_path) / (1024 * 1024)
    if size_mb > max_mb:
        return False, f"File too large: {size_mb:.1f}MB (limit {max_mb}MB)"
    return True, "OK"


def validate_prompt_length(prompt: str) -> Tuple[bool, str]:
    """Validates user input length against upper bounds to prevent context window blowouts."""
    if len(prompt) > AppConfig.MAX_PROMPT_CHARS:
        return False, (
            f"Prompt is too long ({len(prompt):,} chars). "
            f"Max is {AppConfig.MAX_PROMPT_CHARS:,}."
        )
    return True, "OK"


def export_chat_to_markdown(messages: List[Dict], summary: str, title: str) -> str:
    """Formats conversation history and summary metadata into Markdown format for export."""
    md = f"# 💬 {title}\n\n"
    md += f"**Exported**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
    if summary:
        md += "## 📝 Conversation Summary\n\n"
        md += f"{summary}\n\n"
    md += "---\n\n## Full Conversation\n\n"
    for msg in messages:
        role = msg.get("role", "unknown").upper()
        content = msg.get("content", "").strip()
        emoji = "👤" if role == "USER" else "🤖"
        md += f"### {emoji} {role}\n\n{content}\n\n---\n\n"
    return md


# ================================================================================
# HISTORY PERSISTENCE MANAGER
# Thread-safe, atomic JSON storage for persistent multi-session chat history.
# Auto-prunes entries older than 7 days.
# ================================================================================

_HISTORY_LOCK = threading.Lock()


class HistoryManager:
    """
    Handles JSON file IO for chat sessions. Features atomic file replacements
    (writing to `.tmp` before renaming) to guarantee data integrity against crashes.
    """

    @staticmethod
    def load_history() -> Dict[str, Any]:
        """Loads and filters stored chat sessions, purging entries older than retention bounds."""
        if not os.path.exists(AppConfig.HISTORY_FILE):
            return {}
        try:
            with open(AppConfig.HISTORY_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
            cutoff = datetime.now() - timedelta(days=AppConfig.HISTORY_RETENTION_DAYS)
            valid_history = {}
            for session_id, session in data.items():
                updated_str = session.get("updated_at")
                if updated_str:
                    try:
                        updated_time = datetime.fromisoformat(updated_str)
                        if updated_time >= cutoff:
                            valid_history[session_id] = session
                    except ValueError:
                        valid_history[session_id] = session
                else:
                    valid_history[session_id] = session
            return valid_history
        except json.JSONDecodeError:
            print("⚠️ Chat history file is corrupted. Resetting to empty history.")
            return {}
        except Exception as e:
            print(f"❌ History load error: {friendly_error(e)}")
            return {}

    @staticmethod
    def save_session(session_id: str, title: str, messages: List[Dict[str, str]],
                     file_metadata: Optional[str] = None,
                     summary: Optional[str] = None):
        """Atomically persists session updates using thread-safe write locks."""
        tmp_path = AppConfig.HISTORY_FILE + ".tmp"
        try:
            with _HISTORY_LOCK:
                history = HistoryManager.load_history()
                now_str = datetime.now().isoformat()
                created_at = history.get(session_id, {}).get("created_at", now_str)
                history[session_id] = {
                    "title": title, "file_metadata": file_metadata,
                    "created_at": created_at, "updated_at": now_str,
                    "summary": summary or "", "messages": messages,
                }
                with open(tmp_path, "w", encoding="utf-8") as f:
                    json.dump(history, f, indent=2, ensure_ascii=False)
                os.replace(tmp_path, AppConfig.HISTORY_FILE)
        except Exception as e:
            print(f"❌ Error saving chat session: {friendly_error(e)}")
            if os.path.exists(tmp_path):
                try:
                    os.remove(tmp_path)
                except Exception:
                    pass

    @staticmethod
    def delete_session(session_id: str):
        """Deletes a designated session ID from disk persistence atomically."""
        tmp_path = AppConfig.HISTORY_FILE + ".tmp"
        try:
            with _HISTORY_LOCK:
                history = HistoryManager.load_history()
                if session_id in history:
                    del history[session_id]
                with open(tmp_path, "w", encoding="utf-8") as f:
                    json.dump(history, f, indent=2, ensure_ascii=False)
                os.replace(tmp_path, AppConfig.HISTORY_FILE)
        except Exception as e:
            print(f"❌ Error deleting chat session: {friendly_error(e)}")
            if os.path.exists(tmp_path):
                try:
                    os.remove(tmp_path)
                except Exception:
                    pass


def get_google_llm(model: str, **kwargs) -> ChatGoogleGenerativeAI:
    """Helper wrapper for fetching cached Google Generative AI LLM instances."""
    temperature = kwargs.pop("temperature", 0.0)
    return _get_cached_chat_llm(model, temperature)


# ================================================================================
# CONVERSATION SUMMARIZER & ADAPTIVE QUERY ENHANCER
# Condenses long dialogue turns and applies advanced retrieval strategies (HyDE, Multi-Query).
# ================================================================================

class ConversationSummarizer:
    """
    Compresses long-running conversations into rolling textual summaries when message limits
    are reached, conserving context window space for core reasoning.
    """
    SUMMARIZER_PROMPT = ChatPromptTemplate.from_template("""
You are a conversation compressor. Condense chat history into a concise summary.

PREVIOUS SUMMARY:
{previous_summary}

NEW MESSAGES:
{new_messages}

RULES:
1. Preserve user preferences, key data facts, metrics, decisions, and file names.
2. Omit greetings and fluff.
3. Maximum length: ~{max_tokens} tokens.

OUTPUT ONLY updated summary text:
""")

    @staticmethod
    def _format_messages(messages: List[Dict[str, str]]) -> str:
        lines = []
        for m in messages:
            role = m.get("role", "unknown").upper()
            content = m.get("content", "").strip()
            if len(content) > 800:
                content = content[:800] + "..."
            lines.append(f"[{role}]: {content}")
        return "\n".join(lines)

    @classmethod
    def maybe_summarize(cls, llm, messages, existing_summary, config: Optional[Dict[str, Any]] = None):
        """Triggers progressive summarization if conversation turn count exceeds trigger limits."""
        n = len(messages)
        if n < AppConfig.SUMMARY_TRIGGER_LENGTH:
            return existing_summary, messages
        split_idx = max(0, n - AppConfig.MAX_VERBATIM_MESSAGES)
        old_messages = messages[:split_idx]
        recent_messages = messages[split_idx:]
        if not old_messages:
            return existing_summary, recent_messages
        try:
            chain = cls.SUMMARIZER_PROMPT | llm | StrOutputParser()
            new_summary = invoke_with_retry(chain, {
                "previous_summary": existing_summary or "(none)",
                "new_messages": cls._format_messages(old_messages),
                "max_tokens": AppConfig.MAX_SUMMARY_TOKENS,
            }, config=config).strip()
            return new_summary, recent_messages
        except Exception:
            return existing_summary, messages


class QueryEnhancer:
    """
    Determines optimal retrieval strategies (Direct, Multi-Query Expansion, Query Decomposition,
    or Hypothetical Document Embeddings - HyDE) to maximize vector retrieval recall and precision.
    """
    def __init__(self, llm, config: Optional[Dict[str, Any]] = None):
        self.llm = llm
        self.config = config or {}

    def _invoke_llm(self, template: str, **kwargs) -> str:
        try:
            prompt = ChatPromptTemplate.from_template(template)
            chain = prompt | self.llm | StrOutputParser()
            return chain.invoke(kwargs, config=self.config).strip()
        except Exception as e:
            raise RuntimeError(f"Query Optimizer error: {e}") from e

    def determine_strategy(self, query: str) -> str:
        """Classifies the prompt into one of four distinct retrieval strategies."""
        template = """
Analyze the request and pick the best retrieval strategy:
- NONE: Specific lookups or direct quotes.
- EXPANSION: Synonyms or industry terminology variations needed.
- DECOMPOSITION: Multi-part or comparative analysis required.
- HYDE: Conceptual questions benefiting from a hypothetical answer draft.

User Request: {query}
Respond ONLY with one word: NONE, EXPANSION, DECOMPOSITION, or HYDE.
"""
        try:
            decision = self._invoke_llm(template, query=query).upper()
            decision = re.sub(r"[^A-Z]", "", decision)
            if decision in ["NONE", "EXPANSION", "DECOMPOSITION", "HYDE"]:
                return decision.lower()
        except Exception:
            pass
        return "hyde"

    def expand_query(self, query: str, num_variations: int = 3) -> List[str]:
        """Generates multiple rephrasings of a user query for multi-query vector expansion."""
        template = """
Generate exactly {num_variations} variations of the query below.
Query: {query}
Return ONLY rephrasings, one per line.
"""
        try:
            raw = self._invoke_llm(template, query=query, num_variations=num_variations)
            variations = [
                re.sub(r"^\s*[\d\.\-\*]+\s*", "", line).strip()
                for line in raw.splitlines() if line.strip()
            ]
            if query not in variations:
                variations.insert(0, query)
            return variations[:num_variations + 1]
        except Exception:
            return [query]

    def decompose_query(self, query: str) -> List[str]:
        """Decomposes a complex prompt into independent sub-questions."""
        template = """
Break down this request into 2-3 simpler sub-questions.
Query: {query}
Return ONLY sub-questions, one per line.
"""
        try:
            raw = self._invoke_llm(template, query=query)
            sub_queries = [
                re.sub(r"^\s*[\d\.\-\*]+\s*", "", line).strip()
                for line in raw.splitlines() if line.strip()
            ]
            return [q for q in sub_queries if q] or [query]
        except Exception:
            return [query]

    def hyde_projection(self, query: str) -> str:
        """Generates a hypothetical answer passage (HyDE) to search for semantically similar chunks."""
        template = """
Draft a short 2-sentence hypothetical answer passage that directly addresses the query.
Query: {query}
Hypothetical Answer:
"""
        try:
            return self._invoke_llm(template, query=query)
        except Exception:
            return query


# ================================================================================
# UNIFIED RESOURCE HANDLER
# Manages multi-modal ingestion, MySQL schema exploration, RAG vector index creation,
# sandboxed Pandas execution, and Gemini Vision reasoning.
# ================================================================================

class ResourceHandler:
    """
    Central controller for processing multi-modal input channels (PDF, DOCX, CSV, Image, MySQL).
    Directs queries to appropriate sub-engines (FAISS/BM25 Hybrid RAG, Sandboxed Code Exec, SQL, Vision).
    """
    def __init__(self):
        self.file_type: Optional[str] = None
        self.active_file_name: Optional[str] = None
        self.vectorstore: Optional[FAISS] = None
        self.hybrid_retriever: Optional[EnsembleRetriever] = None
        self.llm = None
        self.vision_llm = None
        self.enhancer: Optional[QueryEnhancer] = None
        self.langsmith_config: Dict[str, Any] = {}
        self.image_base64: Optional[str] = None
        self.image_mime_type: Optional[str] = None
        self.df: Optional[pd.DataFrame] = None
        self.df_metadata_summary: Optional[str] = None
        self.db_engine = None
        self.selected_db: Optional[str] = None
        self.schema_context: Optional[str] = None
        self.discovered_databases: Dict[str, List[str]] = {}

    def inject_dependencies(self, llm, vision_llm, config: Optional[Dict[str, Any]] = None) -> None:
        """Injects model instances and observability configuration into the resource handler."""
        self.llm = llm
        self.vision_llm = vision_llm
        self.langsmith_config = config or {}
        self.enhancer = QueryEnhancer(llm, config=self.langsmith_config)

    # ---------------------------------------------------------------- MySQL Engine ---

    def connect_mysql_server(self, host: str, port: str, user: str, pass_word: str) -> str:
        """
        Establishes connection to a MySQL server, discovers accessible database schemas and
        tables, and prepares schema metadata context for LLM SQL generation.
        """
        host, port, user, pass_word = host.strip(), port.strip(), user.strip(), pass_word.strip()

        if not port.isdigit():
            return f"❌ Invalid Port: '{port}' must be numeric (e.g. 3306)"

        if not host or host.isdigit() or "@" in host or "/" in host:
            return "❌ Invalid Host: provide a hostname or IP like 'localhost' or '127.0.0.1'"

        safe_user = quote_plus(user)
        safe_pass = quote_plus(pass_word)
        conn_url = f"mysql+pymysql://{safe_user}:{safe_pass}@{host}:{port}/"

        try:
            # Operational connection test
            try:
                test_conn = pymysql.connect(
                    host=host, port=int(port), user=user,
                    password=pass_word, connect_timeout=5,
                )
                test_conn.close()
            except pymysql.err.OperationalError as e:
                error_code = e.args[0] if e.args else None
                error_msg = str(e.args[1]) if len(e.args) > 1 else str(e)
                if error_code == 1045:
                    return f"❌ Access Denied (1045): Wrong password for user `{user}`@{host}"
                elif error_code == 2003:
                    return f"❌ Can't Connect (2003): MySQL not running on `{host}:{port}`"
                elif error_code == 1130:
                    return f"❌ Host Not Allowed (1130): User `{user}` not permitted from `{host}`"
                return f"❌ MySQL {error_code}: {error_msg}"

            self.db_engine = create_engine(
                conn_url, pool_pre_ping=True, pool_recycle=1800,
            )
            databases_info: Dict[str, List[str]] = {}

            # Reflect schemas across user-accessible databases
            with self.db_engine.connect() as conn:
                databases = [
                    row[0] for row in conn.execute(text("SHOW DATABASES;")).fetchall()
                ]
                if not databases:
                    return "❌ Connected, but no accessible databases found."

                system_dbs = {'information_schema', 'mysql', 'performance_schema', 'sys'}
                user_databases = [db for db in databases if db not in system_dbs]
                for db in user_databases:
                    try:
                        tables_query = text(f"SHOW TABLES FROM `{db}`;")
                        tables = [row[0] for row in conn.execute(tables_query).fetchall()]
                        databases_info[db] = tables
                    except Exception:
                        databases_info[db] = ["(unable to list tables)"]

            self.discovered_databases = databases_info
            schema_parts = ["AVAILABLE DATABASES AND TABLES:\n",
                            "(Use fully-qualified table names: `db`.`table`)\n"]
            for db_name, tables in databases_info.items():
                schema_parts.append(f"Database: `{db_name}`")
                if tables and tables != ["(unable to list tables)"]:
                    schema_parts.append(
                        f"  Tables ({len(tables)}): " + ", ".join(f"`{t}`" for t in tables)
                    )
                else:
                    schema_parts.append("  Tables: (none accessible)")
                schema_parts.append("")

            self.schema_context = "\n".join(schema_parts)
            self.file_type = "mysql"
            self.active_file_name = f"MySQL Server ({len(databases_info)} DBs)"
            total_tables = sum(
                len(t) for t in databases_info.values()
                if t and t != ["(unable to list tables)"]
            )
            return f"✅ Connected to MySQL server. Found {len(databases_info)} databases and {total_tables} total tables."

        except Exception as e:
            return f"❌ Connection failed: {friendly_error(e)}"

    def classify_mysql_query(self, question: str) -> str:
        """Classifies a database query intent into: display, schema, calculation, or general."""
        q = question.lower().strip()

        display_patterns = [
            r"\bshow\s+(me\s+)?(the\s+)?(all\s+)?(\w+\s+)?tables?\b",
            r"\bdisplay\s+(the\s+)?(\w+\s+)?tables?\b",
            r"\blist\s+(all\s+)?(\w+\s+)?(tables?|columns?|databases|records?)\b",
            r"\bview\s+(the\s+)?(\w+\s+)?tables?\b",
            r"\bget\s+(the\s+)?(\w+\s+)?tables?\b",
            r"\bshow\s+(\w+\s+)?(data|records|rows)\b",
        ]
        if any(re.search(p, q) for p in display_patterns):
            return "display"

        schema_patterns = [
            r"\bwhat\s+(tables|columns|fields)\b",
            r"\bshow\s+(schema|structure|columns)\b",
            r"\bdescribe\s+\w+",
            r"\bhow\s+many\s+(tables|columns)\b",
        ]
        if any(re.search(p, q) for p in schema_patterns):
            return "schema"

        calc_keywords = [
            "count", "sum", "average", "avg", "max", "min", "total",
            "calculate", "compute", "how many", "how much", "revenue",
            "profit", "sales", "group by", "join", "highest", "lowest",
            "top", "bottom", "most", "least", "best", "worst",
            "compare", "trend", "monthly", "yearly", "weekly", "daily",
            "quarterly", "growth", "increase", "decrease", "rank",
            "percentage", "ratio", "statistic", "stat", "metric",
            "metrics",
        ]
        if any(kw in q for kw in calc_keywords):
            return "calculation"

        general_patterns = [
            r"\bweather\b", r"\bforecast\b",
            r"\bjoke\b", r"\brecipe\b", r"\briddle\b",
            r"\b(movie|song|album)\s+name\b",
            r"\bstock\s+price\b", r"\bcryptocurrency\b", r"\bbitcoin\b",
            r"\bwho\s+(is|was)\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b",
        ]
        if any(re.search(p, q) for p in general_patterns):
            return "general"

        return "calculation"

    def _validate_table_ref(self, target_table: str) -> Optional[str]:
        """Validates that a table reference exists within the discovered database schemas."""
        if not target_table or "." not in target_table:
            return None
        db_part, tbl_part = target_table.split(".", 1)
        db_part = db_part.strip("`").strip()
        tbl_part = tbl_part.strip("`").strip()
        if not db_part or not tbl_part:
            return None
        if db_part not in self.discovered_databases:
            return None
        if tbl_part not in self.discovered_databases[db_part]:
            return None
        return f"`{db_part}`.`{tbl_part}`"

    def _extract_sql_table_refs(self, sql_query: str) -> List[Tuple[str, str]]:
        """Parses raw SQL text using regex to extract referenced db.table pairs."""
        refs: List[Tuple[str, str]] = []
        for m in re.finditer(r"`([^`]+)`\.`([^`]+)`", sql_query):
            refs.append((m.group(1), m.group(2)))
        kw_prefix = r"(?i)(?:FROM|JOIN|UPDATE|INTO|TABLE)\s+"
        pattern = rf"{kw_prefix}`?([A-Za-z_][\w-]*)`?\.`?([A-Za-z_][\w-]*)`?"
        for m in re.finditer(pattern, sql_query):
            refs.append((m.group(1), m.group(2)))
        return refs

    def _looks_like_aggregation(self, sql_query: str) -> bool:
        """Determines if SQL query contains aggregation keywords to prevent automatic LIMIT truncation."""
        upper = sql_query.upper()
        return any(
            token in upper
            for token in ("COUNT(", "SUM(", "AVG(", "MIN(", "MAX(",
                          "GROUP BY", "HAVING")
        )

    def display_table_direct(self, question: str) -> str:
        """Fetches direct previews of specified tables without invoking full LLM SQL generation."""
        if not self.discovered_databases:
            return "❌ No database connection active."

        q = question.lower().strip()

        target_db = None
        for db_name in self.discovered_databases.keys():
            if db_name and re.search(
                rf"(^|\b){re.escape(db_name.lower())}($|\b)", q
            ):
                target_db = db_name
                break

        all_tables: List[str] = []
        if target_db:
            tables = self.discovered_databases.get(target_db, [])
            if tables and tables != ["(unable to list tables)"]:
                all_tables.extend(f"{target_db}.{t}" for t in tables)
        else:
            for db_name, tables in self.discovered_databases.items():
                if tables and tables != ["(unable to list tables)"]:
                    all_tables.extend(f"{db_name}.{t}" for t in tables)

        target_table = None
        for tbl in all_tables:
            tbl_name = tbl.split(".")[-1]
            if tbl_name.lower() in {"the", "all", "a", "an", "my", "this", "that", "some"}:
                continue
            if re.search(rf"(^|\b){re.escape(tbl_name.lower())}($|\b)", q):
                target_table = tbl
                break

        if not target_table:
            return self._list_all_tables_direct()

        safe_fqn = self._validate_table_ref(target_table)
        if not safe_fqn:
            return f"❌ Table `{target_table}` was not found in schema."

        try:
            with self.db_engine.connect() as conn:
                df = pd.read_sql(
                    text(f"SELECT * FROM {safe_fqn} LIMIT :lim"),
                    con=conn, params={"lim": 50},
                )
                count = conn.execute(text(f"SELECT COUNT(*) FROM {safe_fqn}")).scalar() or 0

            if df.empty:
                return f"📋 Table `{target_table}` is empty."

            markdown_table = df.to_string(index=False)
            total_note = (
                f"\n\n_Showing 50 of {count} rows. "
                "Use WHERE clauses for specific data._"
                if count > 50 else ""
            )
            return f"📋 Table: `{target_table}`{total_note}\n\n{markdown_table}"
        except Exception as e:
            return f"❌ Error querying table: {friendly_error(e)}"

    def _list_all_tables_direct(self) -> str:
        """Returns a formatted list of all discovered databases and tables."""
        if not self.discovered_databases:
            return "❌ No database connection active."
        output = ["📚 Available Tables:\n"]
        for db_name, tables in self.discovered_databases.items():
            output.append(f"### 🗄️ `{db_name}`")
            if tables and tables != ["(unable to list tables)"]:
                for tbl in tables:
                    output.append(f"- `{db_name}.{tbl}`")
            else:
                output.append("- _(no accessible tables)_")
            output.append("")
        output.append("💡 Tip: Ask 'show me <table> table' or 'how many rows in <table>'.")
        return "\n".join(output)

    def query_mysql(self, question: str):
        """Entry point for handling MySQL interaction based on classified query intent."""
        if not self.db_engine or not self.schema_context:
            return "❌ No active MySQL server connection."
        query_type = self.classify_mysql_query(question)
        if query_type == "display":
            return self.display_table_direct(question)
        if query_type == "schema":
            return f"📊 Schema:\n{self.schema_context}"
        if query_type == "general":
            return None  # Triggers agent fallback to web search
        return self._query_mysql_with_llm(question)

    def _query_mysql_with_llm(self, question: str) -> str:
        """Generates, safety-validates, and executes SELECT queries against the MySQL server."""
        try:
            sql_prompt = ChatPromptTemplate.from_template(
                "You are an expert MySQL Data Analyst.\n"
                "Write a valid MySQL SELECT query for this request.\n\n"
                "Schema Context:\n{schema_context}\n\n"
                "Question: {question}\n\n"
                "RULES:\n"
                "1. Output ONLY executable SQL inside ```sql ... ```.\n"
                "2. Use fully qualified names: `db`.`table`.\n"
                "3. Only write SELECT queries.\n"
                "4. Add LIMIT 100 max unless aggregated.\n"
            )
            chain = sql_prompt | self.llm | StrOutputParser()
            raw_response = invoke_with_retry(
                chain,
                {"schema_context": self.schema_context, "question": question},
                config={**self.langsmith_config, "run_name": "mysql_sql_generator"},
            )

            sql_query = raw_response.strip()
            if "```sql" in sql_query:
                sql_query = sql_query.split("```sql", 1)[1].split("```", 1)[0].strip()
            elif "```" in sql_query:
                sql_query = sql_query.split("```", 1)[1].split("```", 1)[0].strip()

            if "LIMIT" not in sql_query.upper() and not self._looks_like_aggregation(sql_query):
                sql_query = sql_query.rstrip(";").strip() + " LIMIT 100;"

            # Safety net: Validate that all referenced tables exist in discovered schema
            for db_part, tbl_part in self._extract_sql_table_refs(sql_query):
                if db_part in self.discovered_databases:
                    if not self._validate_table_ref(f"{db_part}.{tbl_part}"):
                        return (
                            f"❌ The generated SQL referenced unknown table "
                            f"`{db_part}.{tbl_part}` — refusing to execute for safety."
                        )

            with self.db_engine.connect() as conn:
                df = pd.read_sql(text(sql_query), con=conn)

            if df.empty:
                result_str = "*Query returned 0 rows.*"
            else:
                display_df = df.head(50)
                total_rows = len(df)
                if total_rows > 50:
                    result_str = (
                        f"{display_df.to_string(index=False)}\n"
                        f"_... showing 50 of {total_rows} rows._"
                    )
                else:
                    result_str = f"{display_df.to_string(index=False)}"

            return (
                f"🗄️ Query:\n{sql_query}\n\n"
                f"📊 Results ({len(df)} rows):\n{result_str}"
            )
        except Exception as err:
            return f"❌ SQL Execution Error: {friendly_error(err)}"

    # --------------------------------------------------------------- Hybrid RAG Retrieval ---

    def _create_hybrid_retriever(self, pages: List[Document]) -> int:
        """
        Chunks documents according to document volume and builds an Ensemble Retriever
        combining sparse keyword search (BM25) and dense semantic vector search (FAISS).
        """
        num_pages = len(pages)
        if num_pages <= 5:
            chunk_size = AppConfig.DEFAULT_CHUNK_SIZE_SMALL
            chunk_overlap = AppConfig.DEFAULT_CHUNK_OVERLAP_SMALL
        elif num_pages <= 15:
            chunk_size = AppConfig.DEFAULT_CHUNK_SIZE_MEDIUM
            chunk_overlap = AppConfig.DEFAULT_CHUNK_OVERLAP_MEDIUM
        else:
            chunk_size = AppConfig.DEFAULT_CHUNK_SIZE_LARGE
            chunk_overlap = AppConfig.DEFAULT_CHUNK_OVERLAP_LARGE

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size, chunk_overlap=chunk_overlap,
        )
        chunks = splitter.split_documents(pages)
        if not chunks:
            raise ValueError("Document produced zero readable text chunks.")

        embeddings = get_embedding_provider()
        self.vectorstore = FAISS.from_documents(chunks, embeddings)
        dense_retriever = self.vectorstore.as_retriever(
            search_kwargs={"k": AppConfig.RETRIEVER_K}
        )
        sparse_retriever = BM25Retriever.from_documents(chunks, k=AppConfig.RETRIEVER_K)
        self.hybrid_retriever = EnsembleRetriever(
            retrievers=[sparse_retriever, dense_retriever],
            weights=AppConfig.ENSEMBLE_WEIGHTS,
        )
        return len(chunks)

    def load_resource(self, source_path: str, filename: str) -> str:
        """Routes file attachment or URL to appropriate parsing and indexing engine."""
        try:
            self.active_file_name = filename
            if source_path.startswith(("http://", "https://")):
                return self._parse_url(source_path)
            ext = os.path.splitext(filename)[1].lower()
            if ext in [".pdf", ".docx", ".doc", ".txt"]:
                return self._parse_structured_doc(source_path, ext)
            elif ext in [".csv", ".xlsx", ".xls"]:
                return self._parse_pandas_dataframe(source_path, ext)
            elif ext in [".png", ".jpg", ".jpeg", ".webp", ".gif"]:
                return self._parse_image_resource(source_path, ext)
            return f"❌ Extension '{ext}' not supported."
        except Exception as e:
            return f"❌ Resource loading failed: {friendly_error(e)}"

    def _parse_url(self, url: str) -> str:
        """Parses web pages via WebBaseLoader and builds hybrid vector index."""
        try:
            loader = WebBaseLoader(url)
            pages = loader.load()
            num_chunks = self._create_hybrid_retriever(pages)
            self.file_type = "url"
            return f"✅ Web page indexed: {num_chunks} chunks ready."
        except Exception as e:
            return f"❌ URL parsing failed: {friendly_error(e)}"

    def _parse_structured_doc(self, file_path: str, ext: str) -> str:
        """Parses PDF/DOCX/TXT files with Docling layout conversion and fallbacks."""
        pages: List[Document] = []
        last_error = None
        if DOCLING_AVAILABLE:
            try:
                converter = DocumentConverter()
                result = converter.convert(file_path)
                pages = [Document(
                    page_content=result.document.export_to_markdown(),
                    metadata={"source": self.active_file_name},
                )]
            except Exception as e:
                last_error = f"docling: {e}"

        if not pages and ext == ".pdf":
            try:
                pages = PyPDFLoader(file_path).load()
            except Exception as e:
                last_error = f"{last_error or ''}; pypdf: {e}"
        elif not pages and ext in (".docx", ".doc"):
            try:
                from docx import Document as _DocxDocument
                docx_obj = _DocxDocument(file_path)
                text_content = "\n".join(
                    p.text for p in docx_obj.paragraphs if p.text.strip()
                )
                if not text_content.strip():
                    raise ValueError("python-docx extracted no text.")
                pages = [Document(
                    page_content=text_content,
                    metadata={"source": self.active_file_name},
                )]
            except Exception as e:
                last_error = f"{last_error or ''}; python-docx: {e}"
        elif not pages and ext == ".txt":
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    pages = [Document(
                        page_content=f.read(),
                        metadata={"source": self.active_file_name},
                    )]
            except Exception as e:
                last_error = f"{last_error or ''}; txt: {e}"

        if not pages:
            raise ValueError(
                f"Failed to extract readable content from document. "
                f"Last parser error: {last_error}"
            )

        num_chunks = self._create_hybrid_retriever(pages)
        self.file_type = ext
        return f"✅ Document indexed into {num_chunks} chunks."

    def _parse_pandas_dataframe(self, file_path: str, ext: str) -> str:
        """Reads CSV/Excel files into a Pandas DataFrame and constructs schema context."""
        if ext == ".csv":
            self.df = pd.read_csv(file_path)
        elif ext == ".xlsx":
            self.df = pd.read_excel(file_path, engine="openpyxl")
        else:
            self.df = pd.read_excel(file_path, engine="xlrd")
        self.df_metadata_summary = (
            f"DataFrame Schema:\n"
            f"- Columns: {list(self.df.columns)}\n"
            f"- Shape: {self.df.shape[0]} rows x {self.df.shape[1]} cols\n"
            f"- Data Types:\n{self.df.dtypes.to_string()}\n"
            f"- First 3 rows:\n{self.df.head(3).to_string()}"
        )
        self.file_type = ext
        return f"✅ Dataset loaded: {self.df.shape[0]} rows ready."

    def _parse_image_resource(self, file_path: str, ext: str) -> str:
        """Encodes image files to Base64 payloads for Gemini multi-modal processing."""
        with open(file_path, "rb") as f:
            self.image_base64 = base64.b64encode(f.read()).decode("utf-8")
        mime_types = {
            ".png": "image/png", ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg", ".webp": "image/webp",
            ".gif": "image/gif",
        }
        self.image_mime_type = mime_types.get(ext, "image/jpeg")
        self.file_type = "image"
        return "✅ Image loaded. Vision reasoning active."

    # ---------------------------------------------------------------- Queries Execution ---

    def query_document(self, question: str) -> Tuple[str, Dict[str, Any]]:
        """Queries the indexed document using adaptive optimization and hybrid retrieval synthesis."""
        if not self.hybrid_retriever or not self.enhancer:
            return "❌ No active document indexed.", {}

        strategy = self.enhancer.determine_strategy(question)
        docs = []
        meta = {"strategy": strategy}

        if strategy == "none":
            docs = self.hybrid_retriever.invoke(question)
        elif strategy == "expansion":
            queries = self.enhancer.expand_query(question)
            for q in queries:
                docs.extend(self.hybrid_retriever.invoke(q)[:2])
        elif strategy == "decomposition":
            sub_qs = self.enhancer.decompose_query(question)
            for sq in sub_qs:
                docs.extend(self.hybrid_retriever.invoke(sq)[:2])
        else:
            proj = self.enhancer.hyde_projection(question)
            docs = self.hybrid_retriever.invoke(proj)[:4]

        if not docs:
            return "❌ No relevant context retrieved.", meta

        context = "\n\n---\n\n".join(d.page_content for d in docs)
        prompt = ChatPromptTemplate.from_template(
            "Answer strictly using context below:\n\n{context}\n\nQuestion: {question}"
        )
        chain = prompt | self.llm | StrOutputParser()
        answer = invoke_with_retry(
            chain,
            {"context": context, "question": question},
            config={**self.langsmith_config, "run_name": "rag_synthesis"},
        )
        return answer, meta

    def query_pandas(self, question: str) -> str:
        """
        Generates Python code to answer data questions on Pandas DataFrames.
        Executes code inside a hardened `_SafeProxy` sandbox with static AST checks
        and a self-healing Reflexion loop (up to 3 re-try attempts).
        """
        if self.df is None:
            return "❌ No active dataset loaded."

        # Step 1: Classify query type (Theory vs Code execution)
        try:
            classify_prompt = ChatPromptTemplate.from_template("""
Decide if the user's question requires executing Python code on the DataFrame `df`.
- Output 'CODE' for: calculations, filtering, math, aggregations, statistics, ranking.
- Output 'THEORY' for: schema summaries, column names, dataset description, counts of fields.

Schema:
{df_context}

Question: {question}
Output ONLY 'CODE' or 'THEORY'.
""")
            classify_chain = classify_prompt | self.llm | StrOutputParser()
            classification = invoke_with_retry(
                classify_chain,
                {"df_context": self.df_metadata_summary, "question": question},
                config={**self.langsmith_config, "run_name": "pandas_query_classifier"},
            ).strip().upper()

            if "THEORY" in classification:
                theory_chain = (
                    ChatPromptTemplate.from_template(
                        "Answer based on this dataset schema:\n{df_context}\n\nQuestion: {question}"
                    )
                    | self.llm | StrOutputParser()
                )
                return invoke_with_retry(
                    theory_chain,
                    {"df_context": self.df_metadata_summary, "question": question},
                    config={**self.langsmith_config, "run_name": "pandas_theory_response"},
                ).strip()
        except Exception:
            pass

        # Step 2: Code Generation
        code_gen_template = """
Write executable Python code to answer the user request on pandas DataFrame `df`.
Assign final answer to variable `result`.
Available builtins: {allowed_builtins}
Output ONLY raw python in ```python ... ``` block.

Schema:
{df_context}

Request: {question}
"""
        def parse_code(text_raw: str) -> str:
            if "```python" in text_raw:
                return text_raw.split("```python", 1)[1].split("```", 1)[0].strip()
            if "```" in text_raw:
                return text_raw.split("```", 1)[1].split("```", 1)[0].strip()
            return text_raw.strip()

        chain = (
            ChatPromptTemplate.from_template(code_gen_template)
            | self.llm
            | StrOutputParser()
            | RunnableLambda(parse_code)
        )
        code = invoke_with_retry(
            chain,
            {
                "df_context": self.df_metadata_summary,
                "question": question,
                "allowed_builtins": ", ".join(AppConfig.SANDBOX_BUILTINS.keys()),
            },
            config={**self.langsmith_config, "run_name": "pandas_code_generator"},
        )

        # Step 3: Sandboxed Execution & Reflexion Repair Loop
        last_err = None
        for cycle in range(AppConfig.MAX_CODE_EXECUTION_ATTEMPTS):
            sandbox_env = {
                "__builtins__": AppConfig.SANDBOX_BUILTINS,
                "df": _SafeProxy(self.df) if not isinstance(self.df, _SafeProxy) else self.df,
                "pd": _SafeProxy(pd),
                "result": None,
            }
            try:
                # Static AST Analysis Pass
                violation = validate_sandbox_code(code)
                if violation:
                    raise ValueError(f"Sandbox validation rejected code: {violation}")
                
                # Execute dynamically in restricted sandbox environment
                exec(code, sandbox_env)
                res = sandbox_env.get("result")
                if res is not None:
                    if isinstance(res, pd.DataFrame):
                        return f"📊 Computation Result:\n{res.head(20).to_string(index=False)}"
                    return f"📊 Computation Result:\n{res}"
                raise ValueError("Variable 'result' was not assigned.")
            except Exception as err:
                last_err = err
                if cycle == AppConfig.MAX_CODE_EXECUTION_ATTEMPTS - 1:
                    return f"❌ Sandbox execution failed after {cycle + 1} attempts: {err}"
                # Feed execution exception back to LLM to self-correct code
                fix_prompt = (
                    ChatPromptTemplate.from_template(
                        "Fix Python code.\nFailed code:\n{failed_code}\n"
                        "Error: {err}\nSchema:\n{schema}\nCorrect Python:"
                    )
                    | self.llm
                    | StrOutputParser()
                    | RunnableLambda(parse_code)
                )
                code = invoke_with_retry(
                    fix_prompt,
                    {"failed_code": code, "err": str(err), "schema": self.df_metadata_summary},
                    config={**self.langsmith_config, "run_name": f"pandas_reflexion_cycle_{cycle + 1}"},
                )
        return f"❌ Data analysis retries exhausted: {last_err}"

    def query_vision_engine(self, question: str) -> str:
        """Sends Base64-encoded image payloads to Gemini Vision model for analysis."""
        if not self.image_base64:
            return "❌ No image loaded."
        if not self.image_mime_type:
            return "❌ Image MIME type not detected — please re-upload the image."
        if not self.vision_llm:
            return "❌ Vision model not initialized for this session."

        try:
            payload = HumanMessage(content=[
                {"type": "text", "text": question},
                {"type": "image_url", "image_url": {
                    "url": f"data:{self.image_mime_type};base64,{self.image_base64}"
                }},
            ])
            response = invoke_with_retry(
                self.vision_llm,
                [payload],
                config={**self.langsmith_config, "run_name": "vision_inference"},
            )
            response_text = extract_message_content(response).strip()
            return response_text or "(Vision model returned no text content)"
        except Exception as e:
            return f"❌ Vision analysis error: {friendly_error(e)}"


# ================================================================================
# LANGGRAPH TOOL DEFINITIONS
# Wrappers defining agent capabilities for document search, data computation, SQL, vision, & web.
# ================================================================================

@tool
def web_search_tool(query: str) -> str:
    """Searches live DuckDuckGo web search."""
    if not DDG_AVAILABLE:
        return "❌ duckduckgo-search package not installed."
    try:
        results = []
        with DDGS() as ddgs:
            for r in ddgs.text(query, max_results=4):
                results.append(f"• **{r.get('title')}**: {r.get('body')} ({r.get('href')})")
        return "\n".join(results) if results else "No results found."
    except Exception as e:
        return f"❌ Web search failed: {friendly_error(e)}"


def build_langgraph_tools(resource_handler: ResourceHandler) -> list:
    """Binds dynamic system tools to the active ResourceHandler instance."""
    @tool
    def document_rag_tool(query: str) -> str:
        """Queries attached PDF/DOCX/TXT or Web page."""
        if resource_handler.file_type in [".pdf", ".docx", ".doc", ".txt", "url"]:
            ans, _ = resource_handler.query_document(query)
            return ans
        return "❌ No active document."

    @tool
    def pandas_dataframe_tool(query: str) -> str:
        """Calculates metrics on loaded CSV/Excel datasets."""
        if resource_handler.file_type in [".csv", ".xlsx", ".xls"]:
            return resource_handler.query_pandas(query)
        return "❌ No active spreadsheet dataset."

    @tool
    def mysql_database_tool(query: str) -> str:
        """Queries connected MySQL server. Returns 'NOT_MYSQL_QUERY' for off-topic."""
        if resource_handler.file_type != "mysql":
            return "❌ No active MySQL connection."
        res = resource_handler.query_mysql(query)
        return res if res is not None else "NOT_MYSQL_QUERY"

    @tool
    def vision_reasoning_tool(query: str) -> str:
        """Inspects attached image assets."""
        if resource_handler.file_type == "image":
            return resource_handler.query_vision_engine(query)
        return "❌ No image uploaded."

    return [
        document_rag_tool,
        pandas_dataframe_tool,
        mysql_database_tool,
        vision_reasoning_tool,
        web_search_tool,
    ]


# ================================================================================
# LANGGRAPH WORKFLOW ENGINE
# Compiles stateful ReAct tool routing graph with state checkpoints and memory.
# ================================================================================

class AgentState(TypedDict):
    """LangGraph agent state object holding chat turn history and summary context."""
    messages: Annotated[List[BaseMessage], add_messages]
    summary: str
    _summary_used: Optional[int]


def _session_id_to_thread(session_id: str) -> str:
    """Generates unique thread identifiers for LangGraph checkpointer memory mapping."""
    return f"workstation::{session_id}::{uuid.uuid4().hex}"


def build_langgraph_workflow(resource_handler: ResourceHandler, summary: str = "",
                             langsmith_config: Optional[Dict[str, Any]] = None):
    """Constructs and compiles the executable LangGraph StateGraph agent."""
    chat_llm = get_google_llm(AppConfig.CHAT_MODEL)
    vision_llm = get_google_llm(AppConfig.VISION_MODEL)
    resource_handler.inject_dependencies(chat_llm, vision_llm, config=langsmith_config or {})

    tools = build_langgraph_tools(resource_handler)
    llm_with_tools = chat_llm.bind_tools(tools)

    cfg = langsmith_config or {}

    def agent_node(state: AgentState) -> dict:
        """Core ReAct decision node determining tool routing or final text generation."""
        active_asset = resource_handler.active_file_name or "None"
        asset_type = resource_handler.file_type or "None"

        sys_msg = (
            "You are an Enterprise Multi-Modal Agent.\n"
            f"Active Asset: '{active_asset}' (Type: {asset_type})\n\n"
            "Tool routing:\n"
            "1. `document_rag_tool` → PDF/DOCX/TXT/URL.\n"
            "2. `pandas_dataframe_tool` → CSV/Excel.\n"
            "3. `mysql_database_tool` → MySQL queries. Returns 'NOT_MYSQL_QUERY' for off-topic.\n"
            "4. `vision_reasoning_tool` → Images.\n"
            "5. `web_search_tool` → Current events / general knowledge / follow-ups when "
            "MySQL returns 'NOT_MYSQL_QUERY'.\n"
        )

        messages: List[BaseMessage] = [SystemMessage(content=sys_msg)]

        # Prepend conversation summary context if present
        if state.get("summary") and not state.get("_summary_used"):
            messages.append(SystemMessage(
                content=f"Summary of older conversation turns:\n{state['summary']}"
            ))
            messages.extend(state["messages"])
            response = invoke_with_retry(
                llm_with_tools, messages,
                config={**cfg, "run_name": "agent_with_summary"},
            )
            return {"messages": [response], "_summary_used": 1}

        messages.extend(state["messages"])
        response = invoke_with_retry(
            llm_with_tools, messages,
            config={**cfg, "run_name": "agent_turn"},
        )
        return {"messages": [response]}

    builder = StateGraph(AgentState)
    builder.add_node("agent", agent_node)
    builder.add_node("tools", ToolNode(tools))
    builder.add_edge(START, "agent")
    
    # Conditional edge: route to 'tools' if tool calls exist, or END if complete
    builder.add_conditional_edges(
        "agent", tools_condition, {"tools": "tools", END: END},
    )
    builder.add_edge("tools", "agent")

    checkpointer = InMemorySaver()
    return builder.compile(checkpointer=checkpointer)


# ================================================================================
# CLI SESSION MANAGER
# Manages multi-session state isolation, active session switching, and deletion in terminal.
# ================================================================================

class CLISessionManager:
    """Manages active chat sessions, switching, and workflow state for the CLI workstation."""
    def __init__(self):
        self.sessions: Dict[str, dict] = {}
        self.current_session_id: Optional[str] = None
        self._initialize_state()

    def _new_session_dict(self, title: str = "New Session") -> dict:
        handler = ResourceHandler()
        ls_cfg = get_langsmith_config(
            run_name="multi_tool_agent_session",
            tags=["multi-tool-agent"],
            metadata={"session_title": title},
        )
        return {
            "title": title,
            "messages": [],
            "resource": handler,
            "workflow": build_langgraph_workflow(handler, "", ls_cfg),
            "file_metadata": None,
            "summary": "",
        }

    def _initialize_state(self):
        """Loads saved chat history from disk persistence into active session objects."""
        try:
            history_db = HistoryManager.load_history()
            for s_id, s_data in history_db.items():
                handler = ResourceHandler()
                ls_cfg = get_langsmith_config(
                    run_name=f"multi_tool_agent_{s_id[:8]}",
                    tags=["multi-tool-agent"],
                    metadata={"session_id": s_id},
                )
                self.sessions[s_id] = {
                    "title": s_data.get("title", "Saved Session"),
                    "messages": s_data.get("messages", []),
                    "resource": handler,
                    "workflow": build_langgraph_workflow(
                        handler, s_data.get("summary", ""), ls_cfg,
                    ),
                    "file_metadata": s_data.get("file_metadata"),
                    "summary": s_data.get("summary", ""),
                }
        except Exception as e:
            print(f"❌ History load error: {friendly_error(e)}")
            self.sessions = {}

        if not self.sessions:
            new_id = str(uuid.uuid4())
            self.sessions[new_id] = self._new_session_dict()
            self.current_session_id = new_id
        else:
            self.current_session_id = next(iter(self.sessions.keys()))

    def get_current_session(self) -> Tuple[str, dict]:
        """Returns the active session tuple (session_id, session_dict)."""
        if not self.current_session_id or self.current_session_id not in self.sessions:
            new_id = str(uuid.uuid4())
            self.sessions[new_id] = self._new_session_dict()
            self.current_session_id = new_id
        return self.current_session_id, self.sessions[self.current_session_id]

    def create_session(self, title: str = "New Session") -> str:
        """Instantiates a new isolated workstation session."""
        new_id = str(uuid.uuid4())
        self.sessions[new_id] = self._new_session_dict(title)
        self.current_session_id = new_id
        print(f"✨ Created new session: [{new_id[:8]}] '{title}'")
        return new_id

    def list_sessions(self):
        """Prints all active and saved workstation sessions to terminal."""
        print("\n" + "=" * 60)
        print("  📜 RECENT WORKSTATION SESSIONS")
        print("=" * 60)
        for idx, (s_id, s_data) in enumerate(self.sessions.items(), start=1):
            active_marker = " ➔ " if s_id == self.current_session_id else "   "
            asset_info = f" ({s_data.get('file_metadata')})" if s_data.get('file_metadata') else ""
            print(f"{active_marker}[{idx}] ID: {s_id[:8]} | Title: {s_data.get('title')}{asset_info}")
        print("=" * 60 + "\n")

    def switch_session(self, target: str):
        """Switches current session pointer by numeric index or string ID."""
        target = target.strip()
        if target.isdigit():
            idx = int(target) - 1
            session_ids = list(self.sessions.keys())
            if 0 <= idx < len(session_ids):
                self.current_session_id = session_ids[idx]
                s_data = self.sessions[self.current_session_id]
                print(f"✅ Switched to session [{session_ids[idx][:8]}] '{s_data.get('title')}'")
                return
        elif target in self.sessions:
            self.current_session_id = target
            print(f"✅ Switched to session [{target[:8]}]")
            return
        print(f"❌ Session '{target}' not found.")

    def delete_session(self, target: str):
        """Deletes specified session by index or ID from memory and disk storage."""
        target = target.strip()
        target_id = None
        if target.isdigit():
            idx = int(target) - 1
            session_ids = list(self.sessions.keys())
            if 0 <= idx < len(session_ids):
                target_id = session_ids[idx]
        elif target in self.sessions:
            target_id = target

        if target_id and target_id in self.sessions:
            del self.sessions[target_id]
            HistoryManager.delete_session(target_id)
            print(f"🗑️ Deleted session [{target_id[:8]}]")
            if self.current_session_id == target_id:
                if self.sessions:
                    self.current_session_id = next(iter(self.sessions.keys()))
                else:
                    self.create_session()
        else:
            print(f"❌ Session '{target}' not found.")


# ================================================================================
# INTERACTIVE CLI LOOP & COMMAND HANDLERS
# Interface routines for handling slash commands (/attach, /mysql, /sessions, etc.)
# ================================================================================

def print_banner():
    """Prints ASCII workstation welcome banner and telemetry status."""
    print("\n" + "═" * 70)
    print("  ✨ MULTI TOOL AGENT — LANGGRAPH CLI WORKSTATION")
    print("═" * 70)
    if is_langsmith_enabled():
        print(f"  🟢 LangSmith Observability: CONNECTED ({AppConfig.LANGCHAIN_PROJECT})")
    else:
        print("  🔴 LangSmith Observability: DISABLED")
    print("═" * 70)
    print("  Type '/help' to see available commands or ask any question directly.")
    print("═" * 70 + "\n")


def print_help():
    """Displays CLI command reference manual."""
    print("\n" + "─" * 60)
    print("  🛠️ AVAILABLE CLI COMMANDS")
    print("─" * 60)
    print("  /attach <filepath_or_url>  Attach PDF, DOCX, CSV, Image, or URL")
    print("  /mysql                    Connect to MySQL server interactively")
    print("  /sessions                 List all active workstation sessions")
    print("  /new [title]              Create a new isolated chat session")
    print("  /switch <id_or_number>    Switch active chat session")
    print("  /delete <id_or_number>    Delete a chat session")
    print("  /summary                  Show conversation summary for active session")
    print("  /export                   Export current session history to Markdown")
    print("  /help                     Show this help screen")
    print("  /exit or /quit            Exit workstation")
    print("─" * 60 + "\n")


def handle_attach_command(mgr: CLISessionManager, arg: str):
    """Processes user `/attach` command for local documents or URLs."""
    if not arg:
        print("❌ Please specify a file path or URL. Example: /attach data.csv")
        return

    s_id, session = mgr.get_current_session()
    handler: ResourceHandler = session["resource"]

    if arg.startswith(("http://", "https://")):
        target_path = arg
        fname = arg
    else:
        target_path = os.path.abspath(arg)
        fname = os.path.basename(target_path)
        valid, msg = validate_file_size(target_path, AppConfig.MAX_FILE_SIZE_MB)
        if not valid:
            print(f"❌ {msg}")
            return

    status = handler.load_resource(target_path, fname)
    print(status)
    if "✅" in status:
        session["file_metadata"] = fname
        if session["title"] == "New Session":
            session["title"] = fname[:18]
        ls_cfg = get_langsmith_config(
            run_name=f"attachment_{fname[:16]}",
            tags=["multi-tool-agent", "attachment"],
            metadata={"file": fname, "session_id": s_id},
        )
        session["workflow"] = build_langgraph_workflow(
            handler, session.get("summary", ""), ls_cfg,
        )
        HistoryManager.save_session(
            s_id, session["title"], session["messages"], fname,
            session.get("summary", ""),
        )


def handle_mysql_command(mgr: CLISessionManager):
    """Interactively prompts user for MySQL database credentials and initiates connection."""
    print("\n🗄️ MySQL Connection Setup:")
    host = input("  Host [localhost]: ").strip() or "localhost"
    port = input("  Port [3306]: ").strip() or "3306"
    user = input("  User [root]: ").strip() or "root"
    pass_word = input("  Password: ").strip()

    s_id, session = mgr.get_current_session()
    handler: ResourceHandler = session["resource"]

    status = handler.connect_mysql_server(host, port, user, pass_word)
    print(f"\n{status}\n")

    if "✅" in status:
        session["file_metadata"] = "MySQL"
        session["title"] = "MySQL Server"
        ls_cfg = get_langsmith_config(
            run_name="mysql_session",
            tags=["multi-tool-agent", "mysql"],
            metadata={"session_id": s_id},
        )
        session["workflow"] = build_langgraph_workflow(
            handler, session.get("summary", ""), ls_cfg,
        )
        HistoryManager.save_session(
            s_id, session["title"], session["messages"], "MySQL",
            session.get("summary", ""),
        )


def process_user_prompt(mgr: CLISessionManager, prompt: str):
    """Executes single dialogue turn through summarizer, graph orchestration, and output persistence."""
    valid, msg = validate_prompt_length(prompt)
    if not valid:
        print(f"❌ {msg}")
        return

    s_id, session = mgr.get_current_session()
    session["messages"].append({"role": "user", "content": prompt})
    if session["title"] == "New Session":
        session["title"] = prompt[:18]

    print("\n🤖 Assistant Thinking & Orchestrating Tools...")
    response_text = ""
    try:
        chat_llm = get_google_llm(AppConfig.CHAT_MODEL)
        turn_ls_cfg = get_langsmith_config(
            run_name=f"turn_{s_id[:8]}_{int(time.time())}",
            tags=["multi-tool-agent", "turn"],
            metadata={
                "session_id": s_id,
                "active_file": session.get("file_metadata"),
                "user_prompt_preview": prompt[:120],
            },
        )
        summary, recent_messages = ConversationSummarizer.maybe_summarize(
            chat_llm, session["messages"],
            session.get("summary", ""),
            config=turn_ls_cfg,
        )
        session["summary"] = summary

        langchain_msgs: List[BaseMessage] = []
        for m in recent_messages:
            if m["role"] == "user":
                langchain_msgs.append(HumanMessage(content=m["content"]))
            elif m["role"] == "assistant":
                langchain_msgs.append(AIMessage(content=m["content"]))

        config = {
            "configurable": {
                "thread_id": _session_id_to_thread(s_id),
            },
            "recursion_limit": AppConfig.RECURSION_LIMIT,
            **({"callbacks": get_langsmith_callbacks()} if is_langsmith_enabled() else {}),
        }

        output = session["workflow"].invoke(
            {
                "messages": langchain_msgs,
                "summary": summary,
                "_summary_used": 0,
            },
            config=config,
        )

        ai_messages = [
            m for m in output.get("messages", [])
            if isinstance(m, AIMessage)
        ]
        for m in reversed(ai_messages):
            txt = extract_message_content(m).strip()
            tool_calls = getattr(m, "tool_calls", None) or []
            if txt and not tool_calls:
                response_text = txt
                break
        if not response_text:
            for m in reversed(ai_messages):
                txt = extract_message_content(m).strip()
                if txt:
                    response_text = txt
                    break

        if not response_text:
            response_text = "⚠️ No response generated."

        print(f"\n{response_text}\n")
    except Exception as err:
        response_text = f"❌ Pipeline Error: {friendly_error(err)}"
        print(f"\n{response_text}\n")

    session["messages"].append(
        {"role": "assistant", "content": response_text}
    )

    HistoryManager.save_session(
        s_id,
        session["title"],
        session["messages"],
        session["file_metadata"],
        session.get("summary", ""),
    )


# ================================================================================
# MAIN ENTRY POINT
# REPL loop handling user commands and prompt routing.
# ================================================================================

def run_cli():
    """Main terminal loop reading input, routing commands, and managing session state."""
    print_banner()
    mgr = CLISessionManager()

    while True:
        try:
            s_id, session = mgr.get_current_session()
            asset = session.get("file_metadata") or "None"
            prompt_label = f"[{session['title'][:12]}|Asset:{asset}] 👤 User > "
            user_input = input(prompt_label).strip()

            if not user_input:
                continue

            if user_input.startswith("/"):
                parts = user_input.split(" ", 1)
                cmd = parts[0].lower()
                arg = parts[1].strip() if len(parts) > 1 else ""

                if cmd in ["/exit", "/quit"]:
                    print("👋 Exiting Multi Tool Workstation. Goodbye!")
                    break
                elif cmd == "/help":
                    print_help()
                elif cmd == "/attach":
                    handle_attach_command(mgr, arg)
                elif cmd == "/mysql":
                    handle_mysql_command(mgr)
                elif cmd == "/sessions":
                    mgr.list_sessions()
                elif cmd == "/new":
                    mgr.create_session(arg or "New Session")
                elif cmd == "/switch":
                    mgr.switch_session(arg)
                elif cmd == "/delete":
                    mgr.delete_session(arg)
                elif cmd == "/summary":
                    print(f"\n📝 Summary: {session.get('summary') or '(No summary generated yet)'}\n")
                elif cmd == "/export":
                    md = export_chat_to_markdown(
                        session["messages"], session.get("summary", ""), session.get("title", "Chat"),
                    )
                    out_path = f"chat_{s_id[:8]}_{int(time.time())}.md"
                    with open(out_path, "w", encoding="utf-8") as f:
                        f.write(md)
                    print(f"📥 Chat exported to '{out_path}'")
                else:
                    print(f"❌ Unknown command '{cmd}'. Type '/help' for options.")
            else:
                process_user_prompt(mgr, user_input)

        except (KeyboardInterrupt, EOFError):
            print("\n👋 Exiting Multi Tool Workstation. Goodbye!")
            break


if __name__ == "__main__":
    run_cli()