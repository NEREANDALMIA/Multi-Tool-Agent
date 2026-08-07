"""
==============================================================================
MULTI TOOL AGENT — LANGGRAPH WORKSTATION (PRODUCTION READY - HARDENED v6)
==============================================================================
Production-grade AI workstation with:
- Multi-modal ingestion (PDF, DOCX, TXT, CSV, Excel, Images incl. GIF, URLs, MySQL)
- LangGraph agent tool routing with checkpointer memory
- Hybrid retrieval (BM25 + FAISS vector search)
- Query optimization (NONE, EXPANSION, DECOMPOSITION, HYDE)
- Auto-summarization & 7-day chat history persistence (atomic file writes)
- Self-healing code execution for data analysis (hardened sandbox wrapper)
- DuckDuckGo live web search integration
- Robust SQL safety net (validated table refs + LIMIT skipping for aggregations)
- LangSmith observability (opt-in via env vars; zero overhead when disabled)
==============================================================================
"""

import os
import re
import ast
import uuid
import json
import base64
import tempfile
import time
import threading
from datetime import datetime, timedelta
from typing import List, Optional, Dict, Any, Union, Tuple, Annotated, TypedDict

# Third-party data processing, HTTP networking, UI, and database libraries
import pandas as pd
import requests
import streamlit as st
import pymysql
from dotenv import load_dotenv
from sqlalchemy import create_engine, text
from urllib.parse import quote_plus

# LangChain Core primitives for messaging, document abstractions, output parsing, prompts, and tools
from langchain_core.messages import (
    BaseMessage, AIMessage, HumanMessage, SystemMessage,
)
from langchain_core.documents import Document
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnableLambda
from langchain_core.tools import tool

# Google Gemini integrations for LLM chat and text embeddings
from langchain_google_genai import (
    ChatGoogleGenerativeAI,
    GoogleGenerativeAIEmbeddings,
)

# Web scraping, PDF loader, vector stores, local embeddings, and text chunking splitters
from langchain_community.document_loaders import WebBaseLoader, PyPDFLoader
from langchain_community.retrievers import BM25Retriever
try:
    from langchain_classic.retrievers.ensemble import EnsembleRetriever
except ImportError:
    from langchain_community.retrievers import EnsembleRetriever

from langchain_community.vectorstores import FAISS
from langchain_ollama import OllamaEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

# LangGraph state graph primitives, message manipulators, tool nodes, and memory checkpointers
from langgraph.graph import StateGraph, START, END
from langgraph.graph.message import add_messages
from langgraph.prebuilt import ToolNode, tools_condition
from langgraph.checkpoint.memory import InMemorySaver

# High-fidelity document converter (Docling) for rich markdown parsing
try:
    from docling.document_converter import DocumentConverter
    DOCLING_AVAILABLE = True
except ImportError:
    DOCLING_AVAILABLE = False

# DuckDuckGo live web search integration
try:
    from duckduckgo_search import DDGS
    DDG_AVAILABLE = True
except ImportError:
    DDG_AVAILABLE = False

# ──────────────────────────────────────────────────────────────────────────
# LangSmith Observability Imports (with graceful fallback)
# ──────────────────────────────────────────────────────────────────────────
try:
    from langchain_core.tracers import LangChainTracer
    from langsmith import Client as LangSmithClient
    LANGSMITH_AVAILABLE = True
except ImportError:
    LANGSMITH_AVAILABLE = False
    LangChainTracer = None
    LangSmithClient = None

# Initialize environment variables from .env file
load_dotenv()


# ================================================================================
# CONFIGURATION
# ================================================================================

class AppConfig:
    """
    Central configuration repository storing model defaults, text chunking bounds,
    retrieval weights, execution timeouts, file limits, and observability configurations.
    """
    # Active Generative AI Models
    CHAT_MODEL: str = os.getenv("CHAT_MODEL", "gemini-2.5-flash")
    VISION_MODEL: str = os.getenv("VISION_MODEL", "gemini-2.5-flash")
    EMBED_MODEL: str = os.getenv("EMBED_MODEL", "nomic-embed-text")
    GOOGLE_EMBED_MODEL: str = os.getenv("GOOGLE_EMBED_MODEL", "models/text-embedding-004")
    GOOGLE_API_KEY: str = os.getenv("GOOGLE_API_KEY", "").strip()

    # Dynamic Chunking Dimensions based on document page count
    DEFAULT_CHUNK_SIZE_SMALL: int = 500
    DEFAULT_CHUNK_OVERLAP_SMALL: int = 50
    DEFAULT_CHUNK_SIZE_MEDIUM: int = 1000
    DEFAULT_CHUNK_OVERLAP_MEDIUM: int = 100
    DEFAULT_CHUNK_SIZE_LARGE: int = 2000
    DEFAULT_CHUNK_OVERLAP_LARGE: int = 200

    # Hybrid Retrieval Settings
    RETRIEVER_K: int = 4
    ENSEMBLE_WEIGHTS: List[float] = [0.3, 0.7]  # [BM25 sparse weight, FAISS dense weight]

    # Timeouts, Reflexion Retries, and Payload Constraints
    OLLAMA_TIMEOUT: float = float(os.getenv("OLLAMA_TIMEOUT", "5.0"))
    MAX_CODE_EXECUTION_ATTEMPTS: int = 3
    MIN_TEXT_LENGTH_FOR_PROCESSING: int = 10
    MAX_FILE_SIZE_MB: float = 50.0

    # Chat History File Persistence
    HISTORY_FILE: str = "chat_history_db.json"
    HISTORY_RETENTION_DAYS: int = 7

    # Progressive Conversation Summarization Settings
    MAX_VERBATIM_MESSAGES: int = 6
    SUMMARY_TRIGGER_LENGTH: int = 10
    MAX_SUMMARY_TOKENS: int = 400

    # Safety Guardrails
    MAX_PROMPT_CHARS: int = 200_000
    RECURSION_LIMIT: int = 12

    # Whitelisted built-in functions for sandboxed Python code evaluation
    SANDBOX_BUILTINS: dict = {
        "len": len, "sum": sum, "min": min, "max": max,
        "abs": abs, "round": round, "list": list, "dict": dict,
        "set": set, "tuple": tuple, "str": str, "int": int,
        "float": float, "bool": bool, "range": range,
        "any": any, "all": all, "sorted": sorted, "zip": zip,
        "enumerate": enumerate,
    }

    # ──────────────────────────────────────────────────────────────────────
    # LangSmith Observability Configuration
    # ──────────────────────────────────────────────────────────────────────
    LANGCHAIN_TRACING_V2: bool = os.getenv("LANGCHAIN_TRACING_V2", "false").lower() == "true"
    LANGCHAIN_API_KEY: str = os.getenv("LANGCHAIN_API_KEY", "").strip()
    LANGCHAIN_PROJECT: str = os.getenv("LANGCHAIN_PROJECT", "multi-tool-agent")
    LANGCHAIN_ENDPOINT: str = os.getenv(
        "LANGCHAIN_ENDPOINT", "https://api.smith.langchain.com"
    )


# ──────────────────────────────────────────────────────────────────────────
# LangSmith Helper Functions
# ──────────────────────────────────────────────────────────────────────────
def is_langsmith_enabled() -> bool:
    """
    Determines whether LangSmith tracing is fully configured and operational.
    Requires: langsmith installed, tracing env var enabled, and a valid API key.
    """
    return (
        LANGSMITH_AVAILABLE
        and AppConfig.LANGCHAIN_TRACING_V2
        and bool(AppConfig.LANGCHAIN_API_KEY)
    )


def get_langsmith_callbacks() -> List:
    """
    Constructs the list of LangChain callback handlers for explicit tracing.
    Returns an empty list if LangSmith is not enabled (callers must handle this).
    """
    if not is_langsmith_enabled():
        return []
    callbacks = []
    try:
        client = LangSmithClient(
            api_url=AppConfig.LANGCHAIN_ENDPOINT,
            api_key=AppConfig.LANGCHAIN_API_KEY,
        )
        tracer = LangChainTracer(
            project_name=AppConfig.LANGCHAIN_PROJECT,
            client=client,
        )
        callbacks.append(tracer)
    except Exception as e:
        # Observability is best-effort — never break the pipeline.
        print(f"⚠️ LangSmith tracer initialization failed: {e}")
    return callbacks


def get_langsmith_config(
    run_name: Optional[str] = None,
    tags: Optional[List[str]] = None,
    metadata: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """
    Builds a Runnable config dict carrying LangSmith callbacks + run metadata.
    Pass it as the `config=` arg to any chain's `.invoke()`.
    """
    config: Dict[str, Any] = {}
    callbacks = get_langsmith_callbacks()
    if callbacks:
        config["callbacks"] = callbacks
    if run_name:
        config["run_name"] = run_name
    if tags:
        config["tags"] = tags
    if metadata:
        config["metadata"] = metadata
    return config


# ================================================================================
# SAFE PROXY  (FIX: blocks sandbox MRO escape)
# ================================================================================

class _SafeProxy:
    """
    Wraps an object (typically a pandas DataFrame, Series, or module) and
    blocks access to dangerous dunder attributes that are commonly used in
    Python sandbox escapes (e.g. `obj.__class__.__mro__[1].__subclasses__()`).

    Allowed dunder attributes are a minimal whitelist — everything data-related
    (column access, iteration, indexing, arithmetic operators). Everything
    related to identity, type inspection, or code execution is blocked.
    """

    # Strict whitelist of allowed dunders for data manipulations
    _ALLOWED_DUNDERS = frozenset({
        "__getitem__", "__setitem__", "__delitem__",
        "__iter__", "__next__", "__len__", "__contains__",
        "__add__", "__sub__", "__mul__", "__truediv__", "__floordiv__", "__mod__",
        "__radd__", "__rsub__", "__rmul__", "__rtruediv__",
        "__neg__", "__pos__", "__abs__",
        "__eq__", "__ne__", "__lt__", "__le__", "__gt__", "__ge__",
        "__hash__", "__bool__",
        "__repr__", "__str__",
        "__getattr__",
    })

    _BLOCKED_DUNDER_PREFIXES = ("_",)

    def __init__(self, obj):
        object.__setattr__(self, "_obj", obj)

    def __getattr__(self, name):
        # Intercept dunder attribute access to enforce whitelist check
        if name.startswith("__") and name.endswith("__"):
            if name in self._ALLOWED_DUNDERS:
                return getattr(object.__getattribute__(self, "_obj"), name)
            raise AttributeError(
                f"Access to dunder '{name}' is blocked for security."
            )
        return getattr(object.__getattribute__(self, "_obj"), name)

    def __setattr__(self, name, value):
        if name.startswith("__") and name.endswith("__") and name != "_obj":
            raise AttributeError("Cannot assign dunder attributes.")
        object.__setattr__(self, name, value)

    def __getitem__(self, key):
        return object.__getattribute__(self, "_obj")[key]

    def __iter__(self):
        return iter(object.__getattribute__(self, "_obj"))

    def __len__(self):
        return len(object.__getattribute__(self, "_obj"))

    def __contains__(self, item):
        return item in object.__getattribute__(self, "_obj")

    def __repr__(self):
        return repr(object.__getattribute__(self, "_obj"))


# ================================================================================
# SANDBOX STATIC VALIDATOR
# ================================================================================

# Blacklist of functions, modules, and builtins forbidden inside the execution sandbox
_SANDBOX_FORBIDDEN_NAMES = frozenset({
    "eval", "exec", "compile", "open", "__import__", "globals", "locals",
    "vars", "getattr", "setattr", "delattr", "input", "help", "memoryview",
    "os", "sys", "subprocess", "shutil", "socket", "ctypes", "importlib",
    "pathlib", "pickle", "builtins",
})

# Permitted dunder attributes during AST walking
_SANDBOX_ALLOWED_DUNDER_ATTRS = frozenset({
    "__len__", "__iter__", "__next__", "__contains__",
    "__add__", "__sub__", "__mul__", "__truediv__", "__floordiv__", "__mod__",
    "__radd__", "__rsub__", "__rmul__", "__rtruediv__",
    "__neg__", "__pos__", "__abs__",
    "__eq__", "__ne__", "__lt__", "__le__", "__gt__", "__ge__",
    "__repr__", "__str__",
})


def validate_sandbox_code(code: str) -> Optional[str]:
    """
    Parse generated code into an AST and reject it if it contains any known sandbox
    escape primitive. Returns None if the code is safe, or an error string
    describing why it was rejected.
    """
    try:
        tree = ast.parse(code, mode="exec")
    except SyntaxError as e:
        return f"Generated code has a syntax error: {e}"

    # Inspect all AST nodes statically prior to dynamic evaluation
    for node in ast.walk(tree):
        if isinstance(node, (ast.Import, ast.ImportFrom)):
            return "Import statements are not allowed in the sandbox."
        if isinstance(node, ast.Attribute):
            attr = node.attr
            if attr.startswith("__") and attr.endswith("__"):
                if attr not in _SANDBOX_ALLOWED_DUNDER_ATTRS:
                    return (
                        f"Access to dunder attribute '{attr}' is blocked "
                        "(sandbox-escape primitive)."
                    )
        if isinstance(node, ast.Name) and node.id in _SANDBOX_FORBIDDEN_NAMES:
            return f"Reference to '{node.id}' is not allowed in the sandbox."

    return None


# ================================================================================
# LLM / EMBEDDING CACHE
# ================================================================================

@st.cache_resource(show_spinner=False)
def _get_cached_chat_llm(model_name: str, temperature: float = 0.0):
    """Retrieves or constructs a cached ChatGoogleGenerativeAI client."""
    if not AppConfig.GOOGLE_API_KEY:
        st.error("⚠️ GOOGLE_API_KEY is missing! Set it in .env or environment variables.")
        st.stop()
    return ChatGoogleGenerativeAI(
        model=model_name,
        google_api_key=AppConfig.GOOGLE_API_KEY,
        temperature=temperature,
        max_retries=2,
        timeout=60,
    )


class _EmbeddingProvider:
    """Adaptive embedding provider using local Ollama if reachable, else falling back to Google Cloud Embeddings."""
    def __init__(self):
        self.mode = "google"
        try:
            r = requests.get(
                "http://localhost:11434/api/tags",
                timeout=AppConfig.OLLAMA_TIMEOUT,
            )
            if r.status_code == 200:
                installed = {
                    m.get("name", "").split(":")[0]
                    for m in r.json().get("models", [])
                }
                if AppConfig.EMBED_MODEL.split(":")[0] in installed:
                    self._client = OllamaEmbeddings(model=AppConfig.EMBED_MODEL)
                    self.mode = "ollama"
                    return
        except requests.RequestException:
            pass

        if not AppConfig.GOOGLE_API_KEY:
            raise ValueError("GOOGLE_API_KEY required for fallback embeddings.")
        self._client = GoogleGenerativeAIEmbeddings(
            model=AppConfig.GOOGLE_EMBED_MODEL,
            google_api_key=AppConfig.GOOGLE_API_KEY,
        )

    @property
    def client(self):
        return self._client


@st.cache_resource(show_spinner=False)
def _get_cached_embeddings():
    """Singleton getter for cached embedding provider."""
    return _EmbeddingProvider()


def get_embedding_provider():
    """Returns initialized embedding provider client."""
    return _get_cached_embeddings().client


# ================================================================================
# UTILITIES
# ================================================================================

def extract_message_content(msg) -> str:
    """Normalizes raw LLM output messages, blocks, or dictionaries into clean text."""
    if msg is None:
        return ""
    content = msg.content if hasattr(msg, "content") else msg

    if isinstance(content, str):
        return content

    if isinstance(content, list):
        text_parts = []
        for block in content:
            if isinstance(block, dict):
                if "text" in block and block["text"]:
                    text_parts.append(block["text"])
                elif block.get("type") == "text" and block.get("text"):
                    text_parts.append(block["text"])
            elif isinstance(block, str):
                text_parts.append(block)
        return "\n".join(text_parts) if text_parts else ""

    if isinstance(content, dict):
        if content.get("type") == "text" and content.get("text"):
            return content["text"]
        if "text" in content and content["text"]:
            return content["text"]
        return str(content)

    return str(content) if content is not None else ""


def invoke_with_retry(chain, inputs, max_retries: int = 3, config: Optional[Dict[str, Any]] = None):
    """
    Executes a runnable chain with exponential backoff for transient errors.
    If a LangSmith config dict is supplied, it's merged with the call.
    """
    last_exception = None
    for attempt in range(max_retries):
        try:
            if isinstance(inputs, list):
                return chain.invoke(inputs, config=config) if config else chain.invoke(inputs)
            if isinstance(inputs, str):
                return chain.invoke(inputs, config=config) if config else chain.invoke(inputs)
            if isinstance(inputs, dict):
                try:
                    return chain.invoke(inputs, config=config) if config else chain.invoke(inputs)
                except TypeError:
                    if "messages" in inputs and isinstance(inputs["messages"], list):
                        return chain.invoke(inputs["messages"], config=config) if config else chain.invoke(inputs["messages"])
                    if "input" in inputs:
                        return chain.invoke(inputs["input"], config=config) if config else chain.invoke(inputs["input"])
                    raise
            return chain.invoke(inputs, config=config) if config else chain.invoke(inputs)
        except Exception as e:
            last_exception = e
            err_lower = str(e).lower()
            is_retryable = any(
                kw in err_lower
                for kw in ("429", "rate", "quota", "timeout", "connection",
                           "temporary", "unavailable")
            )
            if is_retryable and attempt < max_retries - 1:
                time.sleep(2 ** attempt)
                continue
            if attempt == max_retries - 1:
                raise RuntimeError(f"Failed after {max_retries} retries: {e}") from e
            raise
    raise RuntimeError("invoke_with_retry exited unexpectedly") from last_exception


def friendly_error(error: Exception) -> str:
    """Converts technical exception stack traces into user-friendly error messages."""
    normalized = str(error).lower().replace(" ", "_").replace("-", "_")
    err_short = str(error)[:500]

    error_map = [
        (r"(?i)\bapi[_-]?key[_-]?invalid\b", "🔑 Invalid API key. Check GOOGLE_API_KEY in .env."),
        (r"(?i)\bpermission[_-]?denied\b", "🔑 Permission denied for this API key."),
        (r"(?i)\bquota[_-]?exceeded\b", "⏱️ API quota exceeded. Please try again later."),
        (r"(?i)\b429\b|\brate[_-]?limit\b", "⏱️ Rate limit hit. Slowing down requests."),
        (r"(?i)\bfile[_-]?too[_-]?large\b", "📦 File exceeds size threshold (max 50MB)."),
        (r"(?i)\bunsupported[_-]?format\b", "📄 Format not supported. Use PDF, DOCX, TXT, CSV, Excel, or Images."),
        (r"(?i)(?<!\w)memory\b|(?i)\boom\b|(?i)out[_\s]?of[_\s]?memory", "💾 Out of memory. Try a smaller dataset or restart the app."),
        (r"(?i)\btimeout\b", "⏱️ Request timed out. Try reducing input size."),
        (r"(?i)\bconnection\b", "🌐 Network connectivity issue."),
        (r"(?i)\bmysql\b|py.?mysql", "🗄️ MySQL connection error. Verify database credentials."),
        (r"(?i)\b404\b|(?i)\bnot[_-]?found\b", "🔍 Resource not found. Check the URL or table name."),
        (r"(?i)\bsandbox\b", "🛡️ Sandbox refused unsafe operation."),
    ]
    for pattern, message in error_map:
        if re.search(pattern, normalized):
            return message
    return f"❌ Error: {err_short}"


def validate_file_size(uploaded_file, max_mb: float = 50.0) -> Tuple[bool, str]:
    """Validates that uploaded file size stays within configured threshold."""
    size_mb = uploaded_file.size / (1024 * 1024)
    if size_mb > max_mb:
        return False, f"File too large: {size_mb:.1f}MB (limit {max_mb}MB)"
    return True, "OK"


def validate_prompt_length(prompt: str) -> Tuple[bool, str]:
    """Validates user input length against maximum character bounds."""
    if len(prompt) > AppConfig.MAX_PROMPT_CHARS:
        return False, (
            f"Prompt is too long ({len(prompt):,} chars). "
            f"Max is {AppConfig.MAX_PROMPT_CHARS:,}."
        )
    return True, "OK"


def export_chat_to_markdown(messages: List[Dict], summary: str, title: str) -> str:
    """Formats active chat session into exportable Markdown structure."""
    md = f"# 💬 {title}\n\n"
    md += f"**Exported**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
    if summary:
        md += "## 📝 Conversation Summary\n\n"
        md += f"{summary}\n\n"
    md += "---\n\n## Full Conversation\n\n"
    for msg in messages:
        role = msg.get("role", "unknown").upper()
        content = msg.get("content", "").strip()
        emoji = "👤" if role == "USER" else "🤖"
        md += f"### {emoji} {role}\n\n{content}\n\n---\n\n"
    return md


# ================================================================================
# HISTORY PERSISTENCE
# ================================================================================

_HISTORY_LOCK = threading.Lock()


class HistoryManager:
    """
    JSON-based chat history with:
    - 7-day automatic pruning
    - Atomic file writes (temp + rename)
    - Thread-safe load+modify+write under a single lock acquisition
    """

    @staticmethod
    def load_history() -> Dict[str, Any]:
        """Loads valid non-expired chat sessions from disk."""
        if not os.path.exists(AppConfig.HISTORY_FILE):
            return {}
        try:
            with open(AppConfig.HISTORY_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
            cutoff = datetime.now() - timedelta(days=AppConfig.HISTORY_RETENTION_DAYS)
            valid_history = {}
            for session_id, session in data.items():
                updated_str = session.get("updated_at")
                if updated_str:
                    try:
                        updated_time = datetime.fromisoformat(updated_str)
                        if updated_time >= cutoff:
                            valid_history[session_id] = session
                    except ValueError:
                        valid_history[session_id] = session
                else:
                    valid_history[session_id] = session
            return valid_history
        except json.JSONDecodeError:
            st.warning("Chat history file is corrupted. Resetting to empty history.")
            return {}
        except Exception as e:
            st.error(f"History load error: {friendly_error(e)}")
            return {}

    @staticmethod
    def save_session(session_id: str, title: str, messages: List[Dict[str, str]],
                     file_metadata: Optional[str] = None,
                     summary: Optional[str] = None):
        """Atomically saves session updates using temporary swap file."""
        tmp_path = AppConfig.HISTORY_FILE + ".tmp"
        try:
            with _HISTORY_LOCK:
                history = HistoryManager.load_history()
                now_str = datetime.now().isoformat()
                created_at = history.get(session_id, {}).get("created_at", now_str)
                history[session_id] = {
                    "title": title, "file_metadata": file_metadata,
                    "created_at": created_at, "updated_at": now_str,
                    "summary": summary or "", "messages": messages,
                }
                with open(tmp_path, "w", encoding="utf-8") as f:
                    json.dump(history, f, indent=2, ensure_ascii=False)
                os.replace(tmp_path, AppConfig.HISTORY_FILE)
        except Exception as e:
            st.error(f"Error saving chat session: {friendly_error(e)}")
            if os.path.exists(tmp_path):
                try:
                    os.remove(tmp_path)
                except Exception:
                    pass

    @staticmethod
    def delete_session(session_id: str):
        """Deletes specified session ID from persistent disk storage."""
        tmp_path = AppConfig.HISTORY_FILE + ".tmp"
        try:
            with _HISTORY_LOCK:
                history = HistoryManager.load_history()
                if session_id in history:
                    del history[session_id]
                with open(tmp_path, "w", encoding="utf-8") as f:
                    json.dump(history, f, indent=2, ensure_ascii=False)
                os.replace(tmp_path, AppConfig.HISTORY_FILE)
        except Exception as e:
            st.error(f"Error deleting chat session: {friendly_error(e)}")
            if os.path.exists(tmp_path):
                try:
                    os.remove(tmp_path)
                except Exception:
                    pass


def get_google_llm(model: str, **kwargs) -> ChatGoogleGenerativeAI:
    """Helper getter for cached Google Gemini Chat client."""
    temperature = kwargs.pop("temperature", 0.0)
    return _get_cached_chat_llm(model, temperature)


# ================================================================================
# CONVERSATION SUMMARIZER & QUERY ENHANCER
# ================================================================================

class ConversationSummarizer:
    """Progressive conversation summarizer condensing older turns when thread length grows."""
    SUMMARIZER_PROMPT = ChatPromptTemplate.from_template("""
You are a conversation compressor. Condense chat history into a concise summary.

PREVIOUS SUMMARY:
{previous_summary}

NEW MESSAGES:
{new_messages}

RULES:
1. Preserve user preferences, key data facts, metrics, decisions, and file names.
2. Omit greetings and fluff.
3. Maximum length: ~{max_tokens} tokens.

OUTPUT ONLY updated summary text:
""")

    @staticmethod
    def _format_messages(messages: List[Dict[str, str]]) -> str:
        lines = []
        for m in messages:
            role = m.get("role", "unknown").upper()
            content = m.get("content", "").strip()
            if len(content) > 800:
                content = content[:800] + "..."
            lines.append(f"[{role}]: {content}")
        return "\n".join(lines)

    @classmethod
    def maybe_summarize(cls, llm, messages, existing_summary, config: Optional[Dict[str, Any]] = None):
        """Condenses conversation turns into rolling text summary when trigger limit is reached."""
        n = len(messages)
        if n < AppConfig.SUMMARY_TRIGGER_LENGTH:
            return existing_summary, messages
        split_idx = max(0, n - AppConfig.MAX_VERBATIM_MESSAGES)
        old_messages = messages[:split_idx]
        recent_messages = messages[split_idx:]
        if not old_messages:
            return existing_summary, recent_messages
        try:
            chain = cls.SUMMARIZER_PROMPT | llm | StrOutputParser()
            new_summary = invoke_with_retry(chain, {
                "previous_summary": existing_summary or "(none)",
                "new_messages": cls._format_messages(old_messages),
                "max_tokens": AppConfig.MAX_SUMMARY_TOKENS,
            }, config=config).strip()
            return new_summary, recent_messages
        except Exception:
            return existing_summary, messages


class QueryEnhancer:
    """Classifies user queries and selects optimal search strategies (Direct, Expansion, Decomposition, HyDE)."""
    def __init__(self, llm, config: Optional[Dict[str, Any]] = None):
        self.llm = llm
        self.config = config or {}

    def _invoke_llm(self, template: str, **kwargs) -> str:
        try:
            prompt = ChatPromptTemplate.from_template(template)
            chain = prompt | self.llm | StrOutputParser()
            return chain.invoke(kwargs, config=self.config).strip()
        except Exception as e:
            raise RuntimeError(f"Query Optimizer error: {e}") from e

    def determine_strategy(self, query: str) -> str:
        """Determines best retrieval technique for the given user question."""
        template = """
Analyze the request and pick the best retrieval strategy:
- NONE: Specific lookups or direct quotes.
- EXPANSION: Synonyms or industry terminology variations needed.
- DECOMPOSITION: Multi-part or comparative analysis required.
- HYDE: Conceptual questions benefiting from a hypothetical answer draft.

User Request: {query}
Respond ONLY with one word: NONE, EXPANSION, DECOMPOSITION, or HYDE.
"""
        try:
            decision = self._invoke_llm(template, query=query).upper()
            decision = re.sub(r"[^A-Z]", "", decision)
            if decision in ["NONE", "EXPANSION", "DECOMPOSITION", "HYDE"]:
                return decision.lower()
        except Exception:
            pass
        return "hyde"

    def expand_query(self, query: str, num_variations: int = 3) -> List[str]:
        """Generates query variations for multi-query retrieval."""
        template = """
Generate exactly {num_variations} variations of the query below.
Query: {query}
Return ONLY rephrasings, one per line.
"""
        try:
            raw = self._invoke_llm(template, query=query, num_variations=num_variations)
            variations = [
                re.sub(r"^\s*[\d\.\-\*]+\s*", "", line).strip()
                for line in raw.splitlines() if line.strip()
            ]
            if query not in variations:
                variations.insert(0, query)
            return variations[:num_variations + 1]
        except Exception:
            return [query]

    def decompose_query(self, query: str) -> List[str]:
        """Decomposes complex questions into independent sub-queries."""
        template = """
Break down this request into 2-3 simpler sub-questions.
Query: {query}
Return ONLY sub-questions, one per line.
"""
        try:
            raw = self._invoke_llm(template, query=query)
            sub_queries = [
                re.sub(r"^\s*[\d\.\-\*]+\s*", "", line).strip()
                for line in raw.splitlines() if line.strip()
            ]
            return [q for q in sub_queries if q] or [query]
        except Exception:
            return [query]

    def hyde_projection(self, query: str) -> str:
        """Generates hypothetical answer document for HyDE vector projection."""
        template = """
Draft a short 2-sentence hypothetical answer passage that directly addresses the query.
Query: {query}
Hypothetical Answer:
"""
        try:
            return self._invoke_llm(template, query=query)
        except Exception:
            return query


# ================================================================================
# RESOURCE HANDLER
# ================================================================================

class ResourceHandler:
    """Central manager handling state, document indexing, MySQL connections, and data executions."""
    def __init__(self):
        self.file_type: Optional[str] = None
        self.active_file_name: Optional[str] = None
        self.vectorstore: Optional[FAISS] = None
        self.hybrid_retriever: Optional[EnsembleRetriever] = None
        self.llm = None
        self.vision_llm = None
        self.enhancer: Optional[QueryEnhancer] = None
        # LangSmith config — set via inject_dependencies; propagates to all chains.
        self.langsmith_config: Dict[str, Any] = {}
        self.image_base64: Optional[str] = None
        self.image_mime_type: Optional[str] = None
        self.df: Optional[pd.DataFrame] = None
        self.df_metadata_summary: Optional[str] = None
        self.db_engine = None
        self.selected_db: Optional[str] = None
        self.schema_context: Optional[str] = None
        self.discovered_databases: Dict[str, List[str]] = {}

    def inject_dependencies(self, llm, vision_llm, config: Optional[Dict[str, Any]] = None) -> None:
        """Binds active LLM clients and LangSmith config to the resource handler."""
        self.llm = llm
        self.vision_llm = vision_llm
        self.langsmith_config = config or {}
        self.enhancer = QueryEnhancer(llm, config=self.langsmith_config)

    # ---------------------------------------------------------------- MySQL ---

    def connect_mysql_server(self, host: str, port: str, user: str, pass_word: str) -> str:
        """Establishes connection to MySQL database, discovers schema structure and tables."""
        host, port, user, pass_word = host.strip(), port.strip(), user.strip(), pass_word.strip()

        if not port.isdigit():
            return f"❌ **Invalid Port**: '{port}' must be numeric (e.g. 3306)"

        if not host or host.isdigit() or "@" in host or "/" in host:
            return "❌ **Invalid Host**: provide a hostname or IP like 'localhost' or '127.0.0.1'"

        safe_user = quote_plus(user)
        safe_pass = quote_plus(pass_word)
        conn_url = f"mysql+pymysql://{safe_user}:{safe_pass}@{host}:{port}/"

        try:
            try:
                test_conn = pymysql.connect(
                    host=host, port=int(port), user=user,
                    password=pass_word, connect_timeout=5,
                )
                test_conn.close()
            except pymysql.err.OperationalError as e:
                error_code = e.args[0] if e.args else None
                error_msg = str(e.args[1]) if len(e.args) > 1 else str(e)
                if error_code == 1045:
                    return f"❌ **Access Denied (1045)**: Wrong password for user `{user}`@{host}"
                elif error_code == 2003:
                    return f"❌ **Can't Connect (2003)**: MySQL not running on `{host}:{port}`"
                elif error_code == 1130:
                    return f"❌ **Host Not Allowed (1130)**: User `{user}` not permitted from `{host}`"
                return f"❌ **MySQL {error_code}**: {error_msg}"

            self.db_engine = create_engine(
                conn_url, pool_pre_ping=True, pool_recycle=1800,
            )
            databases_info: Dict[str, List[str]] = {}

            with self.db_engine.connect() as conn:
                databases = [
                    row[0] for row in conn.execute(text("SHOW DATABASES;")).fetchall()
                ]
                if not databases:
                    return "❌ Connected, but no accessible databases found."

                system_dbs = {'information_schema', 'mysql', 'performance_schema', 'sys'}
                user_databases = [db for db in databases if db not in system_dbs]
                for db in user_databases:
                    try:
                        tables_query = text(f"SHOW TABLES FROM `{db}`;")
                        tables = [row[0] for row in conn.execute(tables_query).fetchall()]
                        databases_info[db] = tables
                    except Exception:
                        databases_info[db] = ["(unable to list tables)"]

            self.discovered_databases = databases_info
            schema_parts = ["AVAILABLE DATABASES AND TABLES:\n",
                            "(Use fully-qualified table names: `db`.`table`)\n"]
            for db_name, tables in databases_info.items():
                schema_parts.append(f"Database: `{db_name}`")
                if tables and tables != ["(unable to list tables)"]:
                    schema_parts.append(
                        f"  Tables ({len(tables)}): " + ", ".join(f"`{t}`" for t in tables)
                    )
                else:
                    schema_parts.append("  Tables: (none accessible)")
                schema_parts.append("")

            self.schema_context = "\n".join(schema_parts)
            self.file_type = "mysql"
            self.active_file_name = f"MySQL Server ({len(databases_info)} DBs)"
            total_tables = sum(
                len(t) for t in databases_info.values()
                if t and t != ["(unable to list tables)"]
            )
            return f"✅ Connected to MySQL server. Found {len(databases_info)} databases and {total_tables} total tables."

        except Exception as e:
            return f"❌ Connection failed: {friendly_error(e)}"

    def classify_mysql_query(self, question: str) -> str:
        """Classifies SQL user questions into display, schema, calculation, or general categories."""
        q = question.lower().strip()

        display_patterns = [
            r"\bshow\s+(me\s+)?(the\s+)?(all\s+)?(\w+\s+)?tables?\b",
            r"\bdisplay\s+(the\s+)?(\w+\s+)?tables?\b",
            r"\blist\s+(all\s+)?(\w+\s+)?(tables?|columns?|databases|records?)\b",
            r"\bview\s+(the\s+)?(\w+\s+)?tables?\b",
            r"\bget\s+(the\s+)?(\w+\s+)?tables?\b",
            r"\bshow\s+(\w+\s+)?(data|records|rows)\b",
        ]
        if any(re.search(p, q) for p in display_patterns):
            return "display"

        schema_patterns = [
            r"\bwhat\s+(tables|columns|fields)\b",
            r"\bshow\s+(schema|structure|columns)\b",
            r"\bdescribe\s+\w+",
            r"\bhow\s+many\s+(tables|columns)\b",
        ]
        if any(re.search(p, q) for p in schema_patterns):
            return "schema"

        calc_keywords = [
            "count", "sum", "average", "avg", "max", "min", "total",
            "calculate", "compute", "how many", "how much", "revenue",
            "profit", "sales", "group by", "join", "highest", "lowest",
            "top", "bottom", "most", "least", "best", "worst",
            "compare", "trend", "monthly", "yearly", "weekly", "daily",
            "quarterly", "growth", "increase", "decrease", "rank",
            "percentage", "ratio", "statistic", "stat", "metric",
            "metrics",
        ]
        if any(kw in q for kw in calc_keywords):
            return "calculation"

        general_patterns = [
            r"\bweather\b", r"\bforecast\b",
            r"\bjoke\b", r"\brecipe\b", r"\briddle\b",
            r"\b(movie|song|album)\s+name\b",
            r"\bstock\s+price\b", r"\bcryptocurrency\b", r"\bbitcoin\b",
            r"\bwho\s+(is|was)\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b",
        ]
        if any(re.search(p, q) for p in general_patterns):
            return "general"

        return "calculation"

    def _validate_table_ref(self, target_table: str) -> Optional[str]:
        """Validates that a table reference exists within discovered MySQL schema."""
        if not target_table or "." not in target_table:
            return None
        db_part, tbl_part = target_table.split(".", 1)
        db_part = db_part.strip("`").strip()
        tbl_part = tbl_part.strip("`").strip()
        if not db_part or not tbl_part:
            return None
        if db_part not in self.discovered_databases:
            return None
        if tbl_part not in self.discovered_databases[db_part]:
            return None
        return f"`{db_part}`.`{tbl_part}`"

    def _extract_sql_table_refs(self, sql_query: str) -> List[Tuple[str, str]]:
        """Parses raw SQL text using regex to extract database and table reference pairs."""
        refs: List[Tuple[str, str]] = []
        for m in re.finditer(r"`([^`]+)`\.`([^`]+)`", sql_query):
            refs.append((m.group(1), m.group(2)))
        kw_prefix = r"(?i)(?:FROM|JOIN|UPDATE|INTO|TABLE)\s+"
        pattern = rf"{kw_prefix}`?([A-Za-z_][\w-]*)`?\.`?([A-Za-z_][\w-]*)`?"
        for m in re.finditer(pattern, sql_query):
            refs.append((m.group(1), m.group(2)))
        return refs

    def _looks_like_aggregation(self, sql_query: str) -> bool:
        """Checks if SQL query contains aggregation operations to avoid auto LIMIT injection."""
        upper = sql_query.upper()
        return any(
            token in upper
            for token in ("COUNT(", "SUM(", "AVG(", "MIN(", "MAX(",
                          "GROUP BY", "HAVING")
        )

    def display_table_direct(self, question: str) -> str:
        """Directly displays database table contents when user asks to preview or view tables."""
        if not self.discovered_databases:
            return "❌ No database connection active."

        q = question.lower().strip()

        target_db = None
        for db_name in self.discovered_databases.keys():
            if db_name and re.search(
                rf"(^|\b){re.escape(db_name.lower())}($|\b)", q
            ):
                target_db = db_name
                break

        all_tables: List[str] = []
        if target_db:
            tables = self.discovered_databases.get(target_db, [])
            if tables and tables != ["(unable to list tables)"]:
                all_tables.extend(f"{target_db}.{t}" for t in tables)
        else:
            for db_name, tables in self.discovered_databases.items():
                if tables and tables != ["(unable to list tables)"]:
                    all_tables.extend(f"{db_name}.{t}" for t in tables)

        target_table = None
        for tbl in all_tables:
            tbl_name = tbl.split(".")[-1]
            if tbl_name.lower() in {"the", "all", "a", "an", "my", "this", "that", "some"}:
                continue
            if re.search(rf"(^|\b){re.escape(tbl_name.lower())}($|\b)", q):
                target_table = tbl
                break

        if not target_table:
            return self._list_all_tables_direct()

        safe_fqn = self._validate_table_ref(target_table)
        if not safe_fqn:
            return f"❌ Table `{target_table}` was not found in schema."

        try:
            with self.db_engine.connect() as conn:
                df = pd.read_sql(
                    text(f"SELECT * FROM {safe_fqn} LIMIT :lim"),
                    con=conn, params={"lim": 50},
                )
                count = conn.execute(text(f"SELECT COUNT(*) FROM {safe_fqn}")).scalar() or 0

            if df.empty:
                return f"📋 Table `{target_table}` is empty."

            markdown_table = df.to_markdown(index=False)
            total_note = (
                f"\n\n_Showing 50 of {count} rows. "
                "Use WHERE clauses for specific data._"
                if count > 50 else ""
            )
            return f"📋 **Table: `{target_table}`**{total_note}\n\n{markdown_table}"
        except Exception as e:
            return f"❌ Error querying table: {friendly_error(e)}"

    def _list_all_tables_direct(self) -> str:
        """Returns formatted overview list of accessible databases and tables."""
        if not self.discovered_databases:
            return "❌ No database connection active."
        output = ["📚 **Available Tables:**\n"]
        for db_name, tables in self.discovered_databases.items():
            output.append(f"### 🗄️ `{db_name}`")
            if tables and tables != ["(unable to list tables)"]:
                for tbl in tables:
                    output.append(f"- `{db_name}.{tbl}`")
            else:
                output.append("- _(no accessible tables)_")
            output.append("")
        output.append("💡 **Tip:** Ask 'show me `<table>` table' or 'how many rows in `<table>`'.")
        return "\n".join(output)

    def query_mysql(self, question: str):
        """Main router for executing MySQL operations based on classified question type."""
        if not self.db_engine or not self.schema_context:
            return "❌ No active MySQL server connection."
        query_type = self.classify_mysql_query(question)
        if query_type == "display":
            return self.display_table_direct(question)
        if query_type == "schema":
            return f"📊 **Schema:**\n```\n{self.schema_context}\n```"
        if query_type == "general":
            return None
        return self._query_mysql_with_llm(question)

    def _query_mysql_with_llm(self, question: str) -> str:
        """Generates SQL via LLM, validates table references, and executes query against server."""
        try:
            sql_prompt = ChatPromptTemplate.from_template(
                "You are an expert MySQL Data Analyst.\n"
                "Write a valid MySQL SELECT query for this request.\n\n"
                "Schema Context:\n{schema_context}\n\n"
                "Question: {question}\n\n"
                "RULES:\n"
                "1. Output ONLY executable SQL inside ```sql ... ```.\n"
                "2. Use fully qualified names: `db`.`table`.\n"
                "3. Only write SELECT queries.\n"
                "4. Add LIMIT 100 max unless aggregated.\n"
            )
            chain = sql_prompt | self.llm | StrOutputParser()
            raw_response = invoke_with_retry(
                chain,
                {"schema_context": self.schema_context, "question": question},
                config={**self.langsmith_config, "run_name": "mysql_sql_generator"},
            )

            sql_query = raw_response.strip()
            if "```sql" in sql_query:
                sql_query = sql_query.split("```sql", 1)[1].split("```", 1)[0].strip()
            elif "```" in sql_query:
                sql_query = sql_query.split("```", 1)[1].split("```", 1)[0].strip()

            if "LIMIT" not in sql_query.upper() and not self._looks_like_aggregation(sql_query):
                sql_query = sql_query.rstrip(";").strip() + " LIMIT 100;"

            for db_part, tbl_part in self._extract_sql_table_refs(sql_query):
                if db_part in self.discovered_databases:
                    if not self._validate_table_ref(f"{db_part}.{tbl_part}"):
                        return (
                            f"❌ The generated SQL referenced unknown table "
                            f"`{db_part}.{tbl_part}` — refusing to execute for safety."
                        )

            with self.db_engine.connect() as conn:
                df = pd.read_sql(text(sql_query), con=conn)

            if df.empty:
                result_str = "*Query returned 0 rows.*"
            else:
                display_df = df.head(50)
                total_rows = len(df)
                if total_rows > 50:
                    result_str = (
                        f"```text\n{display_df.to_string(index=False)}\n```\n"
                        f"_... showing 50 of {total_rows} rows._"
                    )
                else:
                    result_str = f"```text\n{display_df.to_string(index=False)}\n```"

            return (
                f"🗄️ **Query:**\n```sql\n{sql_query}\n```\n\n"
                f"📊 **Results ({len(df)} rows):**\n{result_str}"
            )
        except Exception as err:
            return f"❌ SQL Execution Error: {friendly_error(err)}"


    # --------------------------------------------------------------- Retrieval ---

    def _create_hybrid_retriever(self, pages: List[Document]) -> int:
        """Splits document pages into dynamic text chunks and initializes BM25 + FAISS hybrid retriever."""
        num_pages = len(pages)
        if num_pages <= 5:
            chunk_size = AppConfig.DEFAULT_CHUNK_SIZE_SMALL
            chunk_overlap = AppConfig.DEFAULT_CHUNK_OVERLAP_SMALL
        elif num_pages <= 15:
            chunk_size = AppConfig.DEFAULT_CHUNK_SIZE_MEDIUM
            chunk_overlap = AppConfig.DEFAULT_CHUNK_OVERLAP_MEDIUM
        else:
            chunk_size = AppConfig.DEFAULT_CHUNK_SIZE_LARGE
            chunk_overlap = AppConfig.DEFAULT_CHUNK_OVERLAP_LARGE

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size, chunk_overlap=chunk_overlap,
        )
        chunks = splitter.split_documents(pages)
        if not chunks:
            raise ValueError("Document produced zero readable text chunks.")

        embeddings = get_embedding_provider()
        self.vectorstore = FAISS.from_documents(chunks, embeddings)
        dense_retriever = self.vectorstore.as_retriever(
            search_kwargs={"k": AppConfig.RETRIEVER_K}
        )
        sparse_retriever = BM25Retriever.from_documents(chunks, k=AppConfig.RETRIEVER_K)
        self.hybrid_retriever = EnsembleRetriever(
            retrievers=[sparse_retriever, dense_retriever],
            weights=AppConfig.ENSEMBLE_WEIGHTS,
        )
        return len(chunks)

    def load_resource(self, source_path: str, filename: str) -> str:
        """Parses incoming files or URLs and routes them to appropriate indexing pipeline."""
        try:
            self.active_file_name = filename
            if source_path.startswith(("http://", "https://")):
                return self._parse_url(source_path)
            ext = os.path.splitext(filename)[1].lower()
            if ext in [".pdf", ".docx", ".doc", ".txt"]:
                return self._parse_structured_doc(source_path, ext)
            elif ext in [".csv", ".xlsx", ".xls"]:
                return self._parse_pandas_dataframe(source_path, ext)
            elif ext in [".png", ".jpg", ".jpeg", ".webp", ".gif"]:
                return self._parse_image_resource(source_path, ext)
            return f"❌ Extension '{ext}' not supported."
        except Exception as e:
            return f"❌ Resource loading failed: {friendly_error(e)}"

    def _parse_url(self, url: str) -> str:
        """Fetches web page content and indexes it into vector retriever."""
        try:
            loader = WebBaseLoader(url)
            pages = loader.load()
            num_chunks = self._create_hybrid_retriever(pages)
            self.file_type = "url"
            return f"✅ Web page indexed: {num_chunks} chunks ready."
        except Exception as e:
            return f"❌ URL parsing failed: {friendly_error(e)}"

    def _parse_structured_doc(self, file_path: str, ext: str) -> str:
        """Loads PDF, Word, or plain text documents using Docling or fallback parsers."""
        pages: List[Document] = []
        last_error = None
        if DOCLING_AVAILABLE:
            try:
                converter = DocumentConverter()
                result = converter.convert(file_path)
                pages = [Document(
                    page_content=result.document.export_to_markdown(),
                    metadata={"source": self.active_file_name},
                )]
            except Exception as e:
                last_error = f"docling: {e}"

        if not pages and ext == ".pdf":
            try:
                pages = PyPDFLoader(file_path).load()
            except Exception as e:
                last_error = f"{last_error or ''}; pypdf: {e}"
        elif not pages and ext in (".docx", ".doc"):
            try:
                from docx import Document as _DocxDocument
                docx_obj = _DocxDocument(file_path)
                text_content = "\n".join(
                    p.text for p in docx_obj.paragraphs if p.text.strip()
                )
                if not text_content.strip():
                    raise ValueError("python-docx extracted no text.")
                pages = [Document(
                    page_content=text_content,
                    metadata={"source": self.active_file_name},
                )]
            except Exception as e:
                last_error = f"{last_error or ''}; python-docx: {e}"
        elif not pages and ext == ".txt":
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    pages = [Document(
                        page_content=f.read(),
                        metadata={"source": self.active_file_name},
                    )]
            except Exception as e:
                last_error = f"{last_error or ''}; txt: {e}"

        if not pages:
            raise ValueError(
                f"Failed to extract readable content from document. "
                f"Last parser error: {last_error}"
            )

        num_chunks = self._create_hybrid_retriever(pages)
        self.file_type = ext
        return f"✅ Document indexed into {num_chunks} chunks."

    def _parse_pandas_dataframe(self, file_path: str, ext: str) -> str:
        """Loads tabular CSV or Excel data into pandas DataFrame and generates metadata context."""
        if ext == ".csv":
            self.df = pd.read_csv(file_path)
        elif ext == ".xlsx":
            self.df = pd.read_excel(file_path, engine="openpyxl")
        else:
            self.df = pd.read_excel(file_path, engine="xlrd")
        self.df_metadata_summary = (
            f"DataFrame Schema:\n"
            f"- Columns: {list(self.df.columns)}\n"
            f"- Shape: {self.df.shape[0]} rows x {self.df.shape[1]} cols\n"
            f"- Data Types:\n{self.df.dtypes.to_string()}\n"
            f"- First 3 rows:\n{self.df.head(3).to_string()}"
        )
        self.file_type = ext
        return f"✅ Dataset loaded: {self.df.shape[0]} rows ready."

    def _parse_image_resource(self, file_path: str, ext: str) -> str:
        """Encodes uploaded images to Base64 payloads for Gemini Vision analysis."""
        with open(file_path, "rb") as f:
            self.image_base64 = base64.b64encode(f.read()).decode("utf-8")
        mime_types = {
            ".png": "image/png", ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg", ".webp": "image/webp",
            ".gif": "image/gif",
        }
        self.image_mime_type = mime_types.get(ext, "image/jpeg")
        self.file_type = "image"
        return "✅ Image loaded. Vision reasoning active."


    # ---------------------------------------------------------------- Queries ---

    def query_document(self, question: str) -> Tuple[str, Dict[str, Any]]:
        """Queries active vector index using adaptive query enhancement strategies."""
        if not self.hybrid_retriever or not self.enhancer:
            return "❌ No active document indexed.", {}

        strategy = self.enhancer.determine_strategy(question)
        docs = []
        meta = {"strategy": strategy}

        if strategy == "none":
            docs = self.hybrid_retriever.invoke(question)
        elif strategy == "expansion":
            queries = self.enhancer.expand_query(question)
            for q in queries:
                docs.extend(self.hybrid_retriever.invoke(q)[:2])
        elif strategy == "decomposition":
            sub_qs = self.enhancer.decompose_query(question)
            for sq in sub_qs:
                docs.extend(self.hybrid_retriever.invoke(sq)[:2])
        else:
            proj = self.enhancer.hyde_projection(question)
            docs = self.hybrid_retriever.invoke(proj)[:4]

        if not docs:
            return "❌ No relevant context retrieved.", meta

        context = "\n\n---\n\n".join(d.page_content for d in docs)
        prompt = ChatPromptTemplate.from_template(
            "Answer strictly using context below:\n\n{context}\n\nQuestion: {question}"
        )
        chain = prompt | self.llm | StrOutputParser()
        answer = invoke_with_retry(
            chain,
            {"context": context, "question": question},
            config={**self.langsmith_config, "run_name": "rag_synthesis"},
        )
        return answer, meta

    def query_pandas(self, question: str) -> str:
        """Generates and executes sandboxed pandas Python code with Reflexion self-correction loop."""
        if self.df is None:
            return "❌ No active dataset loaded."

        try:
            classify_prompt = ChatPromptTemplate.from_template("""
Decide if the user's question requires executing Python code on the DataFrame `df`.
- Output 'CODE' for: calculations, filtering, math, aggregations, statistics, ranking.
- Output 'THEORY' for: schema summaries, column names, dataset description, counts of fields.

Schema:
{df_context}

Question: {question}
Output ONLY 'CODE' or 'THEORY'.
""")
            classify_chain = classify_prompt | self.llm | StrOutputParser()
            classification = invoke_with_retry(
                classify_chain,
                {"df_context": self.df_metadata_summary, "question": question},
                config={**self.langsmith_config, "run_name": "pandas_query_classifier"},
            ).strip().upper()

            if "THEORY" in classification:
                theory_chain = (
                    ChatPromptTemplate.from_template(
                        "Answer based on this dataset schema:\n{df_context}\n\nQuestion: {question}"
                    )
                    | self.llm | StrOutputParser()
                )
                return invoke_with_retry(
                    theory_chain,
                    {"df_context": self.df_metadata_summary, "question": question},
                    config={**self.langsmith_config, "run_name": "pandas_theory_response"},
                ).strip()
        except Exception:
            pass

        code_gen_template = """
Write executable Python code to answer the user request on pandas DataFrame `df`.
Assign final answer to variable `result`.
Available builtins: {allowed_builtins}
Output ONLY raw python in ```python ... ``` block.

Schema:
{df_context}

Request: {question}
"""
        def parse_code(text_raw: str) -> str:
            if "```python" in text_raw:
                return text_raw.split("```python", 1)[1].split("```", 1)[0].strip()
            if "```" in text_raw:
                return text_raw.split("```", 1)[1].split("```", 1)[0].strip()
            return text_raw.strip()

        chain = (
            ChatPromptTemplate.from_template(code_gen_template)
            | self.llm
            | StrOutputParser()
            | RunnableLambda(parse_code)
        )
        code = invoke_with_retry(
            chain,
            {
                "df_context": self.df_metadata_summary,
                "question": question,
                "allowed_builtins": ", ".join(AppConfig.SANDBOX_BUILTINS.keys()),
            },
            config={**self.langsmith_config, "run_name": "pandas_code_generator"},
        )

        last_err = None
        for cycle in range(AppConfig.MAX_CODE_EXECUTION_ATTEMPTS):
            sandbox_env = {
                "__builtins__": AppConfig.SANDBOX_BUILTINS,
                "df": _SafeProxy(self.df) if not isinstance(self.df, _SafeProxy) else self.df,
                "pd": _SafeProxy(pd),
                "result": None,
            }
            try:
                violation = validate_sandbox_code(code)
                if violation:
                    raise ValueError(f"Sandbox validation rejected code: {violation}")
                exec(code, sandbox_env)
                res = sandbox_env.get("result")
                if res is not None:
                    if isinstance(res, pd.DataFrame):
                        return f"📊 **Computation Result:**\n```text\n{res.head(20).to_string(index=False)}\n```"
                    return f"📊 **Computation Result:**\n```text\n{res}\n```"
                raise ValueError("Variable 'result' was not assigned.")
            except Exception as err:
                last_err = err
                if cycle == AppConfig.MAX_CODE_EXECUTION_ATTEMPTS - 1:
                    return f"❌ Sandbox execution failed after {cycle + 1} attempts: {err}"
                fix_prompt = (
                    ChatPromptTemplate.from_template(
                        "Fix Python code.\nFailed code:\n{failed_code}\n"
                        "Error: {err}\nSchema:\n{schema}\nCorrect Python:"
                    )
                    | self.llm
                    | StrOutputParser()
                    | RunnableLambda(parse_code)
                )
                code = invoke_with_retry(
                    fix_prompt,
                    {"failed_code": code, "err": str(err), "schema": self.df_metadata_summary},
                    config={**self.langsmith_config, "run_name": f"pandas_reflexion_cycle_{cycle + 1}"},
                )
        return f"❌ Data analysis retries exhausted: {last_err}"

    def query_vision_engine(self, question: str) -> str:
        """Sends multi-modal vision prompt with Base64 image payload to Gemini Vision."""
        if not self.image_base64:
            return "❌ No image loaded."
        if not self.image_mime_type:
            return "❌ Image MIME type not detected — please re-upload the image."
        if not self.vision_llm:
            return "❌ Vision model not initialized for this session."

        try:
            payload = HumanMessage(content=[
                {"type": "text", "text": question},
                {"type": "image_url", "image_url": {
                    "url": f"data:{self.image_mime_type};base64,{self.image_base64}"
                }},
            ])
            response = invoke_with_retry(
                self.vision_llm,
                [payload],
                config={**self.langsmith_config, "run_name": "vision_inference"},
            )
            response_text = extract_message_content(response).strip()
            return response_text or "(Vision model returned no text content)"
        except Exception as e:
            return f"❌ Vision analysis error: {friendly_error(e)}"


# ================================================================================
# LANGGRAPH TOOLS
# ================================================================================

@tool
def web_search_tool(query: str) -> str:
    """Searches live DuckDuckGo web search."""
    if not DDG_AVAILABLE:
        return "❌ duckduckgo-search package not installed."
    try:
        results = []
        with DDGS() as ddgs:
            for r in ddgs.text(query, max_results=4):
                results.append(f"• **{r.get('title')}**: {r.get('body')} ({r.get('href')})")
        return "\n".join(results) if results else "No results found."
    except Exception as e:
        return f"❌ Web search failed: {friendly_error(e)}"


def build_langgraph_tools(resource_handler: ResourceHandler) -> list:
    """Constructs tool objects bound to current ResourceHandler instance."""
    @tool
    def document_rag_tool(query: str) -> str:
        """Queries attached PDF/DOCX/TXT or Web page."""
        if resource_handler.file_type in [".pdf", ".docx", ".doc", ".txt", "url"]:
            ans, _ = resource_handler.query_document(query)
            return ans
        return "❌ No active document."

    @tool
    def pandas_dataframe_tool(query: str) -> str:
        """Calculates metrics on loaded CSV/Excel datasets."""
        if resource_handler.file_type in [".csv", ".xlsx", ".xls"]:
            return resource_handler.query_pandas(query)
        return "❌ No active spreadsheet dataset."

    @tool
    def mysql_database_tool(query: str) -> str:
        """Queries connected MySQL server. Returns 'NOT_MYSQL_QUERY' for off-topic."""
        if resource_handler.file_type != "mysql":
            return "❌ No active MySQL connection."
        res = resource_handler.query_mysql(query)
        return res if res is not None else "NOT_MYSQL_QUERY"

    @tool
    def vision_reasoning_tool(query: str) -> str:
        """Inspects attached image assets."""
        if resource_handler.file_type == "image":
            return resource_handler.query_vision_engine(query)
        return "❌ No image uploaded."

    return [
        document_rag_tool,
        pandas_dataframe_tool,
        mysql_database_tool,
        vision_reasoning_tool,
        web_search_tool,
    ]


# ================================================================================
# LANGGRAPH WORKFLOW
# ================================================================================

class AgentState(TypedDict):
    """TypedDict defining state fields passed between LangGraph nodes."""
    messages: Annotated[List[BaseMessage], add_messages]
    summary: str
    _summary_used: Optional[int]


def _session_id_to_thread(session_id: str) -> str:
    """Generates unique thread key for checkpointer memory lookup."""
    return f"workstation::{session_id}::{uuid.uuid4().hex}"


def build_langgraph_workflow(resource_handler: ResourceHandler, summary: str = "",
                             langsmith_config: Optional[Dict[str, Any]] = None):
    """Compiles executable LangGraph state workflow for tool orchestration."""
    chat_llm = get_google_llm(AppConfig.CHAT_MODEL)
    vision_llm = get_google_llm(AppConfig.VISION_MODEL)
    resource_handler.inject_dependencies(chat_llm, vision_llm, config=langsmith_config or {})

    tools = build_langgraph_tools(resource_handler)
    llm_with_tools = chat_llm.bind_tools(tools)

    cfg = langsmith_config or {}

    def agent_node(state: AgentState) -> dict:
        """ReAct agent node inspecting state, system prompt, and tools for next action."""
        active_asset = resource_handler.active_file_name or "None"
        asset_type = resource_handler.file_type or "None"

        sys_msg = (
            "You are an Enterprise Multi-Modal Agent.\n"
            f"Active Asset: '{active_asset}' (Type: {asset_type})\n\n"
            "Tool routing:\n"
            "1. `document_rag_tool` → PDF/DOCX/TXT/URL.\n"
            "2. `pandas_dataframe_tool` → CSV/Excel.\n"
            "3. `mysql_database_tool` → MySQL queries. Returns 'NOT_MYSQL_QUERY' for off-topic.\n"
            "4. `vision_reasoning_tool` → Images.\n"
            "5. `web_search_tool` → Current events / general knowledge / follow-ups when "
            "MySQL returns 'NOT_MYSQL_QUERY'.\n"
        )

        messages: List[BaseMessage] = [SystemMessage(content=sys_msg)]

        if state.get("summary") and not state.get("_summary_used"):
            messages.append(SystemMessage(
                content=f"Summary of older conversation turns:\n{state['summary']}"
            ))
            messages.extend(state["messages"])
            response = invoke_with_retry(
                llm_with_tools, messages,
                config={**cfg, "run_name": "agent_with_summary"},
            )
            return {"messages": [response], "_summary_used": 1}

        messages.extend(state["messages"])
        response = invoke_with_retry(
            llm_with_tools, messages,
            config={**cfg, "run_name": "agent_turn"},
        )
        return {"messages": [response]}

    builder = StateGraph(AgentState)
    builder.add_node("agent", agent_node)
    builder.add_node("tools", ToolNode(tools))
    builder.add_edge(START, "agent")
    builder.add_conditional_edges(
        "agent", tools_condition, {"tools": "tools", END: END},
    )
    builder.add_edge("tools", "agent")

    checkpointer = InMemorySaver()
    return builder.compile(checkpointer=checkpointer)


# ================================================================================
# STREAMLIT STATE
# ================================================================================

def _new_session_dict(title: str = "New Session") -> dict:
    """Instantiates a fresh isolated workstation session state dictionary."""
    handler = ResourceHandler()
    ls_cfg = get_langsmith_config(
        run_name="multi_tool_agent_session",
        tags=["multi-tool-agent"],
        metadata={"session_title": title},
    )
    return {
        "title": title,
        "messages": [],
        "resource": handler,
        "workflow": build_langgraph_workflow(handler, "", ls_cfg),
        "file_metadata": None,
        "summary": "",
    }


def initialize_system_state():
    """Initializes Streamlit session storage and loads existing chat sessions from disk persistence."""
    if "sessions" not in st.session_state:
        st.session_state.sessions = {}
    if "current_session_id" not in st.session_state:
        st.session_state.current_session_id = None

    if not st.session_state.sessions:
        try:
            history_db = HistoryManager.load_history()
            for s_id, s_data in history_db.items():
                handler = ResourceHandler()
                ls_cfg = get_langsmith_config(
                    run_name=f"multi_tool_agent_{s_id[:8]}",
                    tags=["multi-tool-agent"],
                    metadata={"session_id": s_id},
                )
                st.session_state.sessions[s_id] = {
                    "title": s_data.get("title", "Saved Session"),
                    "messages": s_data.get("messages", []),
                    "resource": handler,
                    "workflow": build_langgraph_workflow(
                        handler, s_data.get("summary", ""), ls_cfg,
                    ),
                    "file_metadata": s_data.get("file_metadata"),
                    "summary": s_data.get("summary", ""),
                }
        except Exception as e:
            st.error(f"History load error: {friendly_error(e)}")
            st.session_state.sessions = {}

    if not st.session_state.sessions:
        new_id = str(uuid.uuid4())
        st.session_state.sessions[new_id] = _new_session_dict()
        st.session_state.current_session_id = new_id

    if (not st.session_state.current_session_id or
        st.session_state.current_session_id not in st.session_state.sessions):
        st.session_state.current_session_id = next(
            iter(st.session_state.sessions.keys()), None
        )

    if not st.session_state.sessions:
        new_id = str(uuid.uuid4())
        st.session_state.sessions[new_id] = _new_session_dict()
        st.session_state.current_session_id = new_id

    curr_id = st.session_state.current_session_id
    if curr_id and curr_id in st.session_state.sessions:
        curr_session = st.session_state.sessions[curr_id]
        st.session_state.curr_messages = curr_session.get("messages", []) or []
        st.session_state.curr_resource = curr_session.get("resource")
        st.session_state.curr_workflow = curr_session.get("workflow")
        st.session_state.curr_summary = curr_session.get("summary", "") or ""
    else:
        st.session_state.curr_messages = []
        st.session_state.curr_summary = ""


# ================================================================================
# UI STYLING
# ================================================================================

def apply_gemini_styles():
    """Injects dark-mode CSS theme styles for Streamlit components."""
    st.markdown("""
        <style>
            .stApp { background-color: #131314 !important; color: #e3e3e3 !important; }
            div[data-testid="stSidebar"] {
                background-color: #1e1f20 !important;
                border-right: 1px solid #2d2e31 !important;
            }
            .gemini-hero-title {
                font-size: 2.8rem; font-weight: 700;
                background: linear-gradient(135deg, #4285f4 0%, #9b51e0 50%, #e91e63 100%);
                -webkit-background-clip: text; -webkit-text-fill-color: transparent;
                text-align: center; margin-top: 1.5rem;
            }
            .gemini-hero-subtitle {
                font-size: 1.1rem; color: #8e918f;
                text-align: center; margin-bottom: 2rem;
            }
            .stButton>button {
                border-radius: 18px !important;
                background-color: #282a2c !important;
                color: #e3e3e3 !important;
                border: 1px solid #3c4043 !important;
            }
            .stButton>button:hover {
                border-color: #a8c7fa !important;
                background-color: #37393b !important;
            }
            .langsmith-status {
                padding: 0.5rem;
                border-radius: 6px;
                font-size: 0.82rem;
                margin-top: 0.5rem;
            }
        </style>
    """, unsafe_allow_html=True)


# ================================================================================
# MAIN UI
# ================================================================================

def _safe_delete_file(path: Optional[str]) -> None:
    """Helper routine to clean up temporary upload files safely."""
    if not path:
        return
    try:
        if os.path.exists(path):
            os.remove(path)
    except Exception:
        pass


def render_workspace():
    """Main rendering entry point building Streamlit sidebar controls and chat workspace."""
    # ─── CHANGED: Page title now reads "Multi Tool Agent" only ───
    st.set_page_config(
        page_title="Multi Tool Agent",
        page_icon="✨",
        layout="centered",
        initial_sidebar_state="expanded",
    )
    apply_gemini_styles()
    initialize_system_state()

    # -------------------------------------------------- Sidebar Interface ------------------
    with st.sidebar:
        st.markdown(
            "<h3 style='color:#a8c7fa;'>✨ Multi Tool Agent</h3>",
            unsafe_allow_html=True,
        )

        if st.button("➕ New Chat", use_container_width=True, type="primary"):
            new_id = str(uuid.uuid4())
            st.session_state.sessions[new_id] = _new_session_dict()
            st.session_state.current_session_id = new_id
            st.rerun()

        curr_id_check = st.session_state.get("current_session_id")
        if curr_id_check and curr_id_check in st.session_state.get("sessions", {}):
            curr_s = st.session_state.sessions[curr_id_check]
            if curr_s.get("messages"):
                md_content = export_chat_to_markdown(
                    curr_s["messages"], curr_s.get("summary", ""),
                    curr_s.get("title", "Chat"),
                )
                st.download_button(
                    label="📥 Export Chat",
                    data=md_content,
                    file_name=f"chat_{datetime.now():%Y%m%d_%H%M}.md",
                    mime="text/markdown",
                    use_container_width=True,
                )

        st.markdown("---")
        st.markdown(
            "<span style='color:#8e918f; font-size:0.8rem;'>RECENT SESSIONS</span>",
            unsafe_allow_html=True,
        )

        sessions = st.session_state.sessions
        curr_id = st.session_state.current_session_id

        for s_id in list(sessions.keys()):
            s_title = sessions[s_id].get("title", "Untitled")
            is_active = (s_id == curr_id)
            col1, col2 = st.columns([0.85, 0.15])
            with col1:
                if st.button(
                    f"💬 {s_title[:18]}", key=f"nav_{s_id}",
                    use_container_width=True,
                    type="primary" if is_active else "secondary",
                ):
                    st.session_state.current_session_id = s_id
                    st.rerun()
            with col2:
                if st.button("🗑️", key=f"del_{s_id}"):
                    st.session_state.sessions.pop(s_id, None)
                    try:
                        HistoryManager.delete_session(s_id)
                    except Exception:
                        pass
                    if is_active:
                        remaining = [k for k in st.session_state.sessions.keys() if k != s_id]
                        if remaining:
                            st.session_state.current_session_id = remaining[0]
                        else:
                            new_id = str(uuid.uuid4())
                            st.session_state.sessions[new_id] = _new_session_dict()
                            st.session_state.current_session_id = new_id
                    st.rerun()

        st.markdown("---")
        st.markdown(
            "<span style='color:#8e918f; font-size:0.8rem;'>ATTACHMENTS</span>",
            unsafe_allow_html=True,
        )
        tab1, tab2 = st.tabs(["📎 Files", "🗄️ MySQL"])

        with tab1:
            up_file = st.file_uploader(
                "Upload",
                type=["pdf", "docx", "txt", "csv", "xlsx", "png", "jpg",
                      "jpeg", "webp", "gif"],
                label_visibility="collapsed",
                key=f"upload_{curr_id}",
            )
            url_in = st.text_input(
                "Or Web URL", placeholder="https://...",
                label_visibility="collapsed",
                key=f"url_{curr_id}",
            )
            if st.button("Process Attachment", use_container_width=True,
                         key=f"proc_{curr_id}"):
                target_path: Optional[str] = None
                fname: Optional[str] = None
                try:
                    if up_file:
                        valid, msg = validate_file_size(up_file, AppConfig.MAX_FILE_SIZE_MB)
                        if not valid:
                            st.error(msg)
                            st.stop()
                        with tempfile.NamedTemporaryFile(
                            delete=False, suffix=f"_{up_file.name}",
                        ) as tmp:
                            tmp.write(up_file.getvalue())
                            target_path = tmp.name
                            fname = up_file.name
                    elif url_in.strip():
                        target_path = url_in.strip()
                        fname = url_in.strip()

                    if target_path and fname:
                        curr = st.session_state.sessions.get(curr_id)
                        if not curr:
                            st.error("No active session.")
                            st.stop()
                        handler = curr.get("resource")
                        if not handler:
                            st.error("Session resource missing.")
                            st.stop()
                        status = handler.load_resource(target_path, fname)
                        if "✅" in status:
                            curr["file_metadata"] = fname
                            curr["title"] = fname[:18]
                            ls_cfg = get_langsmith_config(
                                run_name=f"attachment_{fname[:16]}",
                                tags=["multi-tool-agent", "attachment"],
                                metadata={"file": fname, "session_id": curr_id},
                            )
                            curr["workflow"] = build_langgraph_workflow(
                                handler, curr.get("summary", ""), ls_cfg,
                            )
                            HistoryManager.save_session(
                                curr_id, curr["title"], curr["messages"], fname,
                                curr.get("summary", ""),
                            )
                            st.toast(status)
                            st.rerun()
                        else:
                            st.error(status)
                finally:
                    if target_path and not (target_path.startswith("http://")
                                             or target_path.startswith("https://")):
                        _safe_delete_file(target_path)

        with tab2:
            h = st.text_input("Host", value="localhost", key=f"db_host_{curr_id}")
            p = st.text_input("Port", value="3306", key=f"db_port_{curr_id}")
            u = st.text_input("User", value="root", key=f"db_user_{curr_id}")
            pw = st.text_input("Password", type="password", key=f"db_pass_{curr_id}")
            if st.button("Connect Server", use_container_width=True,
                         key=f"db_connect_{curr_id}"):
                curr = st.session_state.sessions.get(curr_id)
                if not curr:
                    st.error("No active session.")
                    st.stop()
                handler = curr.get("resource")
                if not handler:
                    st.error("Session resource missing.")
                    st.stop()
                status = handler.connect_mysql_server(h, p, u, pw)
                if "✅" in status:
                    curr["file_metadata"] = "MySQL"
                    curr["title"] = "MySQL Server"
                    ls_cfg = get_langsmith_config(
                        run_name="mysql_session",
                        tags=["multi-tool-agent", "mysql"],
                        metadata={"session_id": curr_id},
                    )
                    curr["workflow"] = build_langgraph_workflow(
                        handler, curr.get("summary", ""), ls_cfg,
                    )
                    HistoryManager.save_session(
                        curr_id, curr["title"], curr["messages"], "MySQL",
                        curr.get("summary", ""),
                    )
                    st.toast(status)
                    st.rerun()
                else:
                    st.error(status)

        if curr_id and curr_id in st.session_state.get("sessions", {}):
            handler = st.session_state.sessions[curr_id].get("resource")
            if handler and handler.file_type == "mysql" and handler.discovered_databases:
                with st.expander(
                    f"📚 Discovered Databases ({len(handler.discovered_databases)})",
                    expanded=False,
                ):
                    for db_name, tables in handler.discovered_databases.items():
                        n = (
                            len(tables)
                            if tables and tables != ["(unable to list tables)"] else 0
                        )
                        st.markdown(f"**🗄️ `{db_name}`** — {n} tables")
                        if tables and tables != ["(unable to list tables)"]:
                            for tbl in tables:
                                st.markdown(f"&nbsp;&nbsp;&nbsp;&nbsp;📋 `{tbl}`")
                        else:
                            st.markdown("&nbsp;&nbsp;&nbsp;&nbsp;_(no accessible tables)_")

        # ─── NEW: LangSmith Status Panel ───
        st.markdown("---")
        st.markdown(
            "<span style='color:#8e918f; font-size:0.8rem;'>OBSERVABILITY</span>",
            unsafe_allow_html=True,
        )
        if is_langsmith_enabled():
            st.markdown(
                f"<div class='langsmith-status' style='background:#064e3b;color:#6ee7b7;'>"
                f"🟢 <b>LangSmith Connected</b><br>"
                f"Project: <code>{AppConfig.LANGCHAIN_PROJECT}</code>"
                f"</div>",
                unsafe_allow_html=True,
            )
        else:
            reason = (
                "package not installed" if not LANGSMITH_AVAILABLE
                else "API key missing" if not AppConfig.LANGCHAIN_API_KEY
                else "LANGCHAIN_TRACING_V2 not 'true'"
            )
            st.markdown(
                f"<div class='langsmith-status' style='background:#3f1d1d;color:#fca5a5;'>"
                f"🔴 <b>LangSmith Disabled</b><br>"
                f"{reason}<br>"
                f"<small>Set <code>LANGCHAIN_TRACING_V2=true</code> in .env to enable.</small>"
                f"</div>",
                unsafe_allow_html=True,
            )

    # -------------------------------------------------- Chat Workspace Layout ------------------
    if not st.session_state.curr_messages:
        st.markdown(
            "<div class='gemini-hero-title'>Where intelligence meets data</div>",
            unsafe_allow_html=True,
        )
        st.markdown(
            "<div class='gemini-hero-subtitle'>Ask queries, perform data analysis, "
            "or search the web seamlessly.</div>",
            unsafe_allow_html=True,
        )

    for msg in st.session_state.curr_messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    if st.session_state.get("curr_summary"):
        with st.expander("📝 Conversation Summary (older turns)", expanded=False):
            st.info(st.session_state.curr_summary)

    prompt = st.chat_input("Ask a question, run analysis, or query database...")
    if prompt:
        valid, msg = validate_prompt_length(prompt)
        if not valid:
            st.error(msg)
            st.stop()

        curr_session = st.session_state.sessions[st.session_state.current_session_id]
        curr_session["messages"].append({"role": "user", "content": prompt})
        if curr_session["title"] == "New Session":
            curr_session["title"] = prompt[:18]

        with st.chat_message("user"):
            st.markdown(prompt)

        with st.chat_message("assistant"):
            with st.spinner("LangGraph orchestrating tools..."):
                response_text = ""
                try:
                    chat_llm = get_google_llm(AppConfig.CHAT_MODEL)
                    # Build a per-turn LangSmith config carrying the user prompt
                    # and session metadata, then propagate it everywhere.
                    turn_ls_cfg = get_langsmith_config(
                        run_name=f"turn_{curr_id[:8]}_{int(time.time())}",
                        tags=["multi-tool-agent", "turn"],
                        metadata={
                            "session_id": curr_id,
                            "active_file": curr_session.get("file_metadata"),
                            "user_prompt_preview": prompt[:120],
                        },
                    )
                    summary, recent_messages = ConversationSummarizer.maybe_summarize(
                        chat_llm, curr_session["messages"],
                        curr_session.get("summary", ""),
                        config=turn_ls_cfg,
                    )
                    curr_session["summary"] = summary

                    langchain_msgs: List[BaseMessage] = []
                    for m in recent_messages:
                        if m["role"] == "user":
                            langchain_msgs.append(HumanMessage(content=m["content"]))
                        elif m["role"] == "assistant":
                            langchain_msgs.append(AIMessage(content=m["content"]))

                    config = {
                        "configurable": {
                            "thread_id": _session_id_to_thread(curr_id),
                        },
                        "recursion_limit": AppConfig.RECURSION_LIMIT,
                        # Merge in LangSmith callbacks so LangGraph node execution is also traced.
                        **({"callbacks": get_langsmith_callbacks()} if is_langsmith_enabled() else {}),
                    }

                    output = curr_session["workflow"].invoke(
                        {
                            "messages": langchain_msgs,
                            "summary": summary,
                            "_summary_used": 0,
                        },
                        config=config,
                    )

                    response_text = ""
                    ai_messages = [
                        m for m in output.get("messages", [])
                        if isinstance(m, AIMessage)
                    ]
                    for m in reversed(ai_messages):
                        txt = extract_message_content(m).strip()
                        tool_calls = getattr(m, "tool_calls", None) or []
                        if txt and not tool_calls:
                            response_text = txt
                            break
                    if not response_text:
                        for m in reversed(ai_messages):
                            txt = extract_message_content(m).strip()
                            if txt:
                                response_text = txt
                                break

                    if not response_text:
                        response_text = "⚠️ No response generated."

                    st.markdown(response_text)
                except Exception as err:
                    response_text = f"❌ Pipeline Error: {friendly_error(err)}"
                    st.error(response_text)

                curr_session["messages"].append(
                    {"role": "assistant", "content": response_text}
                )

        HistoryManager.save_session(
            st.session_state.current_session_id,
            curr_session["title"],
            curr_session["messages"],
            curr_session["file_metadata"],
            curr_session.get("summary", ""),
        )


if __name__ == "__main__":
    render_workspace()